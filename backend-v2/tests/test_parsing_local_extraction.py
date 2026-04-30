from pathlib import Path

from app.api.v1.endpoints.parsing import _decode_zip_filename
from app.services.parsing_service import ParsingService, _extract_text_from_file


def test_local_basic_fields_are_normalized_to_frontend_labels():
    text = """
项目名称：咪咕公司与XX公司网络阵地运营管理系统一期工程
招标编号：MGYD20260500019
采购方式：公开招标
预算金额：120万元
获取招标文件时间：2026年03月01日 09:00 至 2026年03月10日 17:00
投标截止时间：2026年03月20日 10:00
投标保证金：人民币20000元，形式为银行保函
付款方式：验收合格后支付合同金额的90%。
★技术分低于60分，不得参与报价分评审
"""
    service = ParsingService()
    merged = service._merge_extraction_results(
        {
            "project_name": "咪咕公司与XX公司网络阵地运营管理系统一期工程",
            "bid_number": "MGYD20260500019",
            "deadline": "2026-03-20 10:00:00",
            "scoring_criteria": "技术分低于60分，不得参与报价分评审",
            "star_items": [{"name": "技术分限制", "content": "技术分低于60分，不得参与报价分评审"}],
            "confidence": {"project_name": 0.9, "bid_number": 0.9, "deadline": 0.95, "scoring_criteria": 0.8},
        },
        {},
        text,
    )

    assert merged["项目名称"]["value"].startswith("咪咕公司")
    assert merged["招标编号"]["value"] == "MGYD20260500019"
    assert merged["投标截止时间"]["value"] == "2026-03-20 10:00:00"
    assert merged["评分重点"]["value"] == "技术分低于60分，不得参与报价分评审"
    assert merged["标书起始时间"]["value"] == "2026-03-01 09:00:00"
    assert merged["标书结束时间"]["value"] == "2026-03-10 17:00:00"
    assert merged["是否有保证金"]["value"] == "是"
    assert merged["保证金金额"]["value"] == "人民币20000元"
    assert merged["星标项列表"][0]["content"] == "技术分低于60分，不得参与报价分评审"


def test_docx_extraction_includes_table_cells(tmp_path: Path):
    from docx import Document

    path = tmp_path / "招标文件.docx"
    doc = Document()
    doc.add_paragraph("项目名称：测试项目")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "评分项"
    table.cell(0, 1).text = "技术方案 30分"
    doc.save(path)

    text, pages = _extract_text_from_file(path, ".docx")

    assert "项目名称：测试项目" in text
    assert "评分项 | 技术方案 30分" in text
    assert pages


def test_decode_zip_filename_prefers_chinese_candidate():
    mojibake = "▓╔╣║╬─╝■.pdf"
    assert _decode_zip_filename(mojibake) == "采购文件.pdf"
