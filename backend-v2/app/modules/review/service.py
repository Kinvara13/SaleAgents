import re
from io import BytesIO
from dataclasses import dataclass
from datetime import datetime
from uuid import uuid4

from docx import Document
from fastapi import HTTPException, status
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.models.review_clause import ReviewClauseRecord
from app.models.review_feedback import ReviewFeedbackRecord
from app.models.review_issue import ReviewIssueRecord
from app.models.review_job import ReviewJob
from app.models.rule_statistics import RuleStatistics
from app.schemas.review import (
    ReviewFeedbackRequest,
    ReviewFeedbackResponse,
    ReviewJobClause,
    ReviewIssueResolveRequest,
    ReviewJobCreateRequest,
    ReviewJobIssue,
    ReviewJobResponse,
    ReviewJobRerunRequest,
)
from app.schemas.workspace import MetricItem, ReviewIssue
from app.services.llm_client import llm_review_client

SEVERITY_RANK = {"P0": 0, "P1": 1, "P2": 2, "P3": 3}


@dataclass(frozen=True)
class RuleDefinition:
    name: str
    title: str
    issue_type: str
    level: str
    detail: str
    suggestion: str
    patterns: tuple[str, ...]
    document: str
    match_mode: str = "any"


RULE_DEFINITIONS = [
    RuleDefinition(
        name="payment_tail_ratio_guardrail",
        title="付款条款与公司回款红线冲突",
        issue_type="付款风险",
        level="P0",
        detail="尾款或终验回款比例偏高，可能显著增加现金流和验收争议风险。",
        suggestion="优先谈判下调终验尾款比例；若无法调整，需提请经营负责人审批豁免。",
        patterns=("终验", "30%"),
        document="付款条款",
        match_mode="all",
    ),
    RuleDefinition(
        name="unlimited_liability_guardrail",
        title="责任承担存在无限责任表述",
        issue_type="责任风险",
        level="P0",
        detail="条款出现无限责任或兜底赔偿表述，超出常规可接受合同边界。",
        suggestion="增加责任上限并限定赔偿范围，避免使用“全部损失”或“无限责任”表述。",
        patterns=("无限责任",),
        document="违约责任条款",
    ),
    RuleDefinition(
        name="one_side_liability_guardrail",
        title="责任分配明显偏向甲方",
        issue_type="责任风险",
        level="P1",
        detail="条款将主要违约、延误或损失责任单边压给乙方，存在不对等风险。",
        suggestion="补充双方责任边界与免责条件，避免“乙方承担全部责任”类单边措辞。",
        patterns=("乙方承担全部责任", "供应商承担全部责任", "全部责任由乙方承担"),
        document="责任分配条款",
    ),
    RuleDefinition(
        name="auto_renewal_guardrail",
        title="合同存在自动续约条款",
        issue_type="续约风险",
        level="P1",
        detail="合同自动续约可能带来价格、服务范围和责任持续生效的管理风险。",
        suggestion="将自动续约改成双方书面确认后续签，补充续约价格和退出条件。",
        patterns=("自动续约", "自动延续"),
        document="期限与续约条款",
    ),
    RuleDefinition(
        name="acceptance_one_side_guardrail",
        title="验收标准存在单方解释风险",
        issue_type="验收风险",
        level="P1",
        detail="验收标准或结论由甲方单方认定，容易引发交付争议和尾款拖延。",
        suggestion="补充客观验收标准、验收时限与默认视为通过机制。",
        patterns=("甲方有权单方认定", "验收标准以甲方解释为准"),
        document="验收条款",
    ),
    RuleDefinition(
        name="ip_assignment_guardrail",
        title="知识产权归属表述需法务确认",
        issue_type="知识产权风险",
        level="P2",
        detail="成果或交付物知识产权直接归甲方所有，可能与公司标准模板冲突。",
        suggestion="区分背景知识产权与交付成果使用权，避免全部权利无条件转移。",
        patterns=("知识产权归甲方所有", "成果归甲方所有"),
        document="知识产权条款",
    ),
]

SEMANTIC_RESPONSIBILITY_TRIGGERS = (
    "按甲方要求",
    "根据甲方要求",
    "根据甲方需要",
    "满足甲方要求",
    "配合甲方",
    "负责相关工作",
    "负责必要工作",
    "完成甲方交办",
)

SEMANTIC_SCOPE_TRIGGERS = (
    "包括但不限于",
    "不限于",
    "其他相关工作",
    "相关配套工作",
    "一切费用由乙方承担",
    "免费提供",
    "无偿提供",
)

SEMANTIC_CHANGE_TRIGGERS = (
    "甲方有权调整",
    "甲方有权变更",
    "甲方可根据需要调整",
    "乙方应无条件配合",
    "乙方应无条件接受",
)

SEMANTIC_COMMITMENT_TRIGGERS = (
    "确保",
    "保证",
    "承诺",
)

SEMANTIC_OUTCOME_TRIGGERS = (
    "通过验收",
    "稳定运行",
    "零故障",
    "不中断",
    "全部要求",
    "所有要求",
    "随叫随到",
    "立即响应",
)


class ReviewService:
    """Contract review orchestration built on deterministic rule checks."""

    def create_job(self, db: Session, payload: ReviewJobCreateRequest) -> ReviewJobResponse:
        clauses = self._split_structured_clauses(payload.contract_text.strip())
        return self._run_job(
            db=db,
            project_id=payload.project_id,
            contract_name=payload.contract_name.strip(),
            contract_type=payload.contract_type.strip() or "采购合同",
            contract_text=payload.contract_text.strip(),
            clauses=clauses,
            trigger=payload.trigger.strip() or "manual",
        )

    def create_job_from_upload(
        self,
        db: Session,
        *,
        filename: str,
        file_bytes: bytes,
        contract_name: str | None,
        contract_type: str,
        project_id: str | None,
        trigger: str,
    ) -> ReviewJobResponse:
        extracted_text = self._extract_text_from_file(filename=filename, file_bytes=file_bytes)
        clauses = self._split_structured_clauses(extracted_text)
        return self._run_job(
            db=db,
            project_id=project_id,
            contract_name=(contract_name or filename).strip(),
            contract_type=contract_type.strip() or "采购合同",
            contract_text=extracted_text,
            clauses=clauses,
            trigger=trigger.strip() or "upload",
        )

    def rerun_job(self, db: Session, job_id: str, payload: ReviewJobRerunRequest) -> ReviewJobResponse:
        job = self._get_job_record(db, job_id)
        contract_text = (payload.contract_text or job.source_text).strip()
        contract_name = (payload.contract_name or job.contract_name).strip()
        contract_type = (payload.contract_type or job.contract_type).strip()

        db.execute(delete(ReviewIssueRecord).where(ReviewIssueRecord.job_id == job_id))
        db.execute(delete(ReviewClauseRecord).where(ReviewClauseRecord.job_id == job_id))
        db.flush()

        clauses = self._split_structured_clauses(contract_text)
        self._apply_job_result(
            db=db,
            job=job,
            contract_name=contract_name or job.contract_name,
            contract_type=contract_type or job.contract_type,
            contract_text=contract_text,
            clauses=clauses,
            trigger="rerun",
        )
        db.commit()
        db.refresh(job)
        return self._to_job_response(job)

    def get_job(self, db: Session, job_id: str) -> ReviewJobResponse:
        return self._to_job_response(self._get_job_record(db, job_id))

    def get_latest_job(self, db: Session) -> ReviewJobResponse:
        job = self._latest_job(db)
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No review job found.",
            )
        return self._to_job_response(job)

    def list_job_issues(self, db: Session, job_id: str) -> list[ReviewJobIssue]:
        self._get_job_record(db, job_id)
        rows = db.scalars(
            select(ReviewIssueRecord)
            .where(ReviewIssueRecord.job_id == job_id)
            .order_by(ReviewIssueRecord.created_at.asc())
        ).all()
        return [ReviewJobIssue.model_validate(row) for row in sorted(rows, key=self._sort_issue_record)]

    def get_job_summary(self, db: Session, job_id: str) -> list[MetricItem]:
        job = self._get_job_record(db, job_id)
        return [MetricItem(**item) for item in job.summary]

    def export_job_report(self, db: Session, job_id: str) -> str:
        job = self._get_job_record(db, job_id)
        issues = self.list_job_issues(db, job_id)
        clauses = self.list_job_clauses(db, job_id)

        lines = [
            f"# 合同审查报告：{job.contract_name}",
            "",
            f"- 合同类型：{job.contract_type}",
            f"- 任务编号：{job.id}",
            f"- 触发方式：{job.trigger}",
            f"- 当前状态：{job.status}",
            f"- 综合风险：{job.overall_risk}",
            f"- 问题总数：{job.issue_count}",
            f"- 高风险问题数：{job.high_risk_issue_count}",
            "",
            "## 摘要",
        ]

        for item in job.summary:
            lines.append(f"- {item['label']}：{item['value']}（{item.get('hint', '')}）")

        lines.extend(["", "## 审查问题"])
        for index, issue in enumerate(issues, start=1):
            lines.extend(
                [
                    f"### {index}. {issue.title}",
                    f"- 等级：{issue.level}",
                    f"- 状态：{issue.status}",
                    f"- 来源：{self._issue_origin(issue.rule_name)}",
                    f"- 位置：{issue.document}",
                    f"- 说明：{issue.detail}",
                    f"- 证据：{issue.evidence or '无'}",
                    f"- 建议：{issue.suggestion or '无'}",
                    f"- 处理备注：{issue.resolution_note or '无'}",
                    "",
                ]
            )

        lines.extend(["## 条款切分"])
        for clause in clauses:
            lines.append(f"- {clause.title} / #{clause.clause_no}：{clause.content}")

        lines.extend(["", "## 建议动作"])
        for item in job.review_actions:
            lines.append(f"- {item}")

        return "\n".join(lines).strip()

    def export_job_report_docx(self, db: Session, job_id: str) -> bytes:
        job = self._get_job_record(db, job_id)
        issues = self.list_job_issues(db, job_id)
        clauses = self.list_job_clauses(db, job_id)

        document = Document()
        document.add_heading(f"合同审查报告：{job.contract_name}", level=0)

        document.add_paragraph(f"合同类型：{job.contract_type}")
        document.add_paragraph(f"任务编号：{job.id}")
        document.add_paragraph(f"触发方式：{job.trigger}")
        document.add_paragraph(f"当前状态：{job.status}")
        document.add_paragraph(f"综合风险：{job.overall_risk}")
        document.add_paragraph(f"问题总数：{job.issue_count}")
        document.add_paragraph(f"高风险问题数：{job.high_risk_issue_count}")

        document.add_heading("摘要", level=1)
        for item in job.summary:
            document.add_paragraph(f"{item['label']}：{item['value']}（{item.get('hint', '')}）")

        document.add_heading("审查问题", level=1)
        for index, issue in enumerate(issues, start=1):
            document.add_heading(f"{index}. {issue.title}", level=2)
            table = document.add_table(rows=8, cols=2)
            table.style = "Table Grid"
            rows_data = [
                ("等级", issue.level),
                ("状态", issue.status),
                ("来源", self._issue_origin(issue.rule_name)),
                ("位置", issue.document),
                ("说明", issue.detail),
                ("证据", issue.evidence or "无"),
                ("建议", issue.suggestion or "无"),
                ("处理备注", issue.resolution_note or "无"),
            ]
            for row_idx, (label, value) in enumerate(rows_data):
                table.rows[row_idx].cells[0].text = label
                table.rows[row_idx].cells[1].text = value

        document.add_heading("条款切分", level=1)
        for clause in clauses:
            document.add_paragraph(f"{clause.title} / #{clause.clause_no}：{clause.content}")

        document.add_heading("建议动作", level=1)
        for item in job.review_actions:
            document.add_paragraph(item, style="List Bullet")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def list_job_clauses(self, db: Session, job_id: str) -> list[ReviewJobClause]:
        self._get_job_record(db, job_id)
        rows = db.scalars(
            select(ReviewClauseRecord)
            .where(ReviewClauseRecord.job_id == job_id)
            .order_by(ReviewClauseRecord.clause_no.asc())
        ).all()
        return [ReviewJobClause.model_validate(row) for row in rows]

    def resolve_issue(
        self, db: Session, issue_id: str, payload: ReviewIssueResolveRequest
    ) -> ReviewJobIssue:
        issue = db.get(ReviewIssueRecord, issue_id)
        if issue is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Review issue '{issue_id}' not found.",
            )

        issue.status = payload.status.strip() or "已处理"
        issue.resolution_note = payload.resolution_note.strip()
        db.commit()
        db.refresh(issue)
        self._refresh_job_summary(db, issue.job_id)
        return ReviewJobIssue.model_validate(issue)

    def submit_feedback(
        self, db: Session, issue_id: str, payload: ReviewFeedbackRequest
    ) -> ReviewFeedbackResponse:
        from uuid import uuid4

        from app.models.review_feedback import ReviewFeedbackRecord
        from app.services.rule_config_service import rule_config_service

        issue = db.get(ReviewIssueRecord, issue_id)
        if issue is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Review issue '{issue_id}' not found.",
            )

        feedback = ReviewFeedbackRecord(
            id=f"feedback-{uuid4().hex[:10]}",
            issue_id=issue_id,
            job_id=issue.job_id,
            rule_name=issue.rule_name,
            feedback_type=payload.feedback_type,
            feedback_note=payload.feedback_note,
            reviewer=payload.reviewer,
        )
        db.add(feedback)
        db.commit()
        db.refresh(feedback)

        rule_config_service.record_feedback(db, issue.rule_name, payload.feedback_type)

        return ReviewFeedbackResponse.model_validate(feedback)

    def get_latest_summary(self, db: Session) -> list[MetricItem] | None:
        job = self._latest_job(db)
        if job is None:
            return None
        return [MetricItem(**item) for item in job.summary]

    def get_latest_issues(self, db: Session) -> list[ReviewIssue] | None:
        job = self._latest_job(db)
        if job is None:
            return None
        issues = self.list_job_issues(db, job.id)
        return [
            ReviewIssue(
                title=item.title,
                type=item.type,
                level=item.level,
                status=item.status,
                document=item.document,
                detail=item.detail,
                evidence=item.evidence,
                suggestion=item.suggestion,
                origin=self._issue_origin(item.rule_name),
                rule_name=item.rule_name,
            )
            for item in issues
        ]

    def get_latest_actions(self, db: Session) -> list[str] | None:
        job = self._latest_job(db)
        if job is None:
            return None
        return list(job.review_actions)

    def _run_job(
        self,
        db: Session,
        project_id: str | None,
        contract_name: str,
        contract_type: str,
        contract_text: str,
        clauses: list[dict[str, str]],
        trigger: str,
    ) -> ReviewJobResponse:
        if not contract_text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Contract text is required.",
            )

        job = ReviewJob(
            id=f"review-job-{uuid4().hex[:10]}",
            project_id=project_id,
            contract_name=contract_name,
            contract_type=contract_type,
            source_text=contract_text,
            trigger=trigger,
            status="running",
        )
        db.add(job)
        db.flush()

        self._apply_job_result(
            db=db,
            job=job,
            contract_name=contract_name,
            contract_type=contract_type,
            contract_text=contract_text,
            clauses=clauses,
            trigger=trigger,
        )
        db.commit()
        db.refresh(job)
        return self._to_job_response(job)

    def _apply_job_result(
        self,
        db: Session,
        job: ReviewJob,
        contract_name: str,
        contract_type: str,
        contract_text: str,
        clauses: list[dict[str, str]],
        trigger: str,
    ) -> None:
        rule_issues = self._evaluate_rules(contract_text, clauses)
        semantic_issues = llm_review_client.review_contract_semantics(
            contract_name=contract_name,
            contract_type=contract_type,
            clauses=clauses,
            rule_issues=rule_issues,
        )
        if not semantic_issues:
            semantic_issues = self._evaluate_semantics(clauses)
        issues = self._merge_issues(rule_issues + semantic_issues)
        review_actions = self._build_actions(issues)
        summary = self._build_summary(issues)
        overall_risk = self._overall_risk(issues)

        job.contract_name = contract_name
        job.contract_type = contract_type
        job.source_text = contract_text
        job.trigger = trigger
        job.status = "done"
        job.overall_risk = overall_risk
        job.issue_count = len(issues)
        job.high_risk_issue_count = len([item for item in issues if item["level"] in {"P0", "P1"}])
        job.summary = [item.model_dump() for item in summary]
        job.review_actions = review_actions
        job.completed_at = datetime.utcnow()

        for clause in clauses:
            db.add(
                ReviewClauseRecord(
                    id=f"review-clause-{uuid4().hex[:10]}",
                    job_id=job.id,
                    clause_no=clause["clause_no"],
                    title=clause["title"],
                    content=clause["content"],
                    source_ref=clause["source_ref"],
                )
            )

        for item in issues:
            db.add(
                ReviewIssueRecord(
                    id=f"review-issue-{uuid4().hex[:10]}",
                    job_id=job.id,
                    title=item["title"],
                    type=item["type"],
                    level=item["level"],
                    status=item["status"],
                    document=item["document"],
                    detail=item["detail"],
                    evidence=item["evidence"],
                    suggestion=item["suggestion"],
                    rule_name=item["rule_name"],
                )
            )

    def _evaluate_rules(self, contract_text: str, clauses: list[dict[str, str]]) -> list[dict[str, str]]:
        normalized_text = self._normalize_text(contract_text)
        matches: list[dict[str, str]] = []
        clause_texts = [item["content"] for item in clauses] or self._split_clauses(contract_text)

        for rule in RULE_DEFINITIONS:
            if not self._matches_rule(rule, normalized_text):
                continue

            evidence = self._find_evidence(clause_texts, rule.patterns) or contract_text[:160].strip()
            matches.append(
                {
                    "title": rule.title,
                    "type": rule.issue_type,
                    "level": rule.level,
                    "status": "待处理",
                    "document": rule.document,
                    "detail": rule.detail,
                    "evidence": evidence,
                    "suggestion": rule.suggestion,
                    "rule_name": rule.name,
                }
            )

        if not matches:
            matches.append(
                {
                    "title": "未命中高风险规则",
                    "type": "规则结论",
                    "level": "P3",
                    "status": "已完成",
                    "document": "规则引擎输出",
                    "detail": "当前规则未发现付款、责任、续约、验收或知识产权方面的显著红线条款。",
                    "evidence": contract_text[:160].strip(),
                    "suggestion": "建议继续补充 LLM 语义审查，检查措辞歧义与隐含责任风险。",
                    "rule_name": "no_rule_hit",
                }
            )

        deduped: dict[str, dict[str, str]] = {}
        for item in matches:
            deduped[item["title"]] = item
        return sorted(deduped.values(), key=self._sort_issue_dict)

    def _evaluate_semantics(self, clauses: list[dict[str, str]]) -> list[dict[str, str]]:
        matches: list[dict[str, str]] = []

        for clause in clauses:
            content = clause["content"]
            normalized = self._normalize_text(content)
            document = self._infer_document_label(content)
            clause_ref = f"{document} · {clause['title']}"

            if self._has_any(normalized, SEMANTIC_RESPONSIBILITY_TRIGGERS):
                matches.append(
                    self._build_semantic_issue(
                        title="责任边界表述模糊，存在兜底履约风险",
                        issue_type="语义责任风险",
                        level="P1",
                        document=clause_ref,
                        evidence=content,
                        detail=(
                            "条款使用“按甲方要求”“配合甲方”等开放式表述，但没有同步限定"
                            "双方配合边界、交付清单或不在范围项，容易把额外协调、返工和兜底义务"
                            "默认压给乙方。"
                        ),
                        suggestion=(
                            "把开放式责任改成明确清单，补充乙方交付范围、甲方配合义务、第三方依赖"
                            "和不属于本次合同的排除项。"
                        ),
                        rule_name="semantic_ambiguous_responsibility",
                    )
                )

            if self._has_any(normalized, SEMANTIC_SCOPE_TRIGGERS):
                matches.append(
                    self._build_semantic_issue(
                        title="服务范围存在隐含扩张风险",
                        issue_type="语义范围风险",
                        level="P1",
                        document=clause_ref,
                        evidence=content,
                        detail=(
                            "条款含有“包括但不限于”“其他相关工作”等兜底措辞，可能把未明示的配套工作、"
                            "额外材料或免费支持一并纳入履约范围，造成范围失控。"
                        ),
                        suggestion=(
                            "将兜底表述拆成可验收的服务清单；新增内容需通过书面变更单确认，并联动"
                            "工期与费用调整。"
                        ),
                        rule_name="semantic_scope_expansion",
                    )
                )

            if self._has_any(normalized, SEMANTIC_CHANGE_TRIGGERS):
                matches.append(
                    self._build_semantic_issue(
                        title="甲方保留单方变更权，缺少费用与工期联动",
                        issue_type="语义变更风险",
                        level="P1",
                        document=clause_ref,
                        evidence=content,
                        detail=(
                            "条款允许甲方单方调整需求、范围或节奏，但没有同步约定变更确认流程，"
                            "也没有把费用、交期和资源追加写清，后续容易出现无偿加项。"
                        ),
                        suggestion=(
                            "补充变更单机制，约定需求范围变化时须书面确认，并同步调整交付计划、"
                            "报价和验收口径。"
                        ),
                        rule_name="semantic_unilateral_change",
                    )
                )

            if self._has_any(normalized, SEMANTIC_COMMITMENT_TRIGGERS) and self._has_any(
                normalized, SEMANTIC_OUTCOME_TRIGGERS
            ):
                matches.append(
                    self._build_semantic_issue(
                        title="结果性承诺缺少前置条件和免责边界",
                        issue_type="语义承诺风险",
                        level="P2",
                        document=clause_ref,
                        evidence=content,
                        detail=(
                            "条款直接承诺“确保/保证”某类结果，但没有写明该结果依赖的输入条件、"
                            "甲方配合义务或第三方约束，容易把合理努力义务放大成绝对结果责任。"
                        ),
                        suggestion=(
                            "将绝对承诺改为在需求冻结、甲方按时配合、第三方接口可用等前提下的"
                            "交付承诺，并补充免责和顺延条件。"
                        ),
                        rule_name="semantic_absolute_commitment",
                    )
                )

        return sorted(self._merge_issues(matches), key=self._sort_issue_dict)

    def _build_summary(self, issues: list[dict[str, str]]) -> list[MetricItem]:
        p0_count = len([item for item in issues if item["level"] == "P0"])
        p1_count = len([item for item in issues if item["level"] == "P1"])
        open_count = len([item for item in issues if item["status"] not in {"已处理", "已关闭"}])
        closed_count = len([item for item in issues if item["status"] in {"已处理", "已关闭"}])
        return [
            MetricItem(label="P0 风险", value=str(p0_count), tone="blue", hint="必须阻断或升级审批"),
            MetricItem(label="P1 风险", value=str(p1_count), tone="cyan", hint="需法务或经营复核"),
            MetricItem(label="待处理项", value=str(open_count), tone="amber", hint="仍需人工确认或修订"),
            MetricItem(label="已关闭", value=str(closed_count), tone="violet", hint="已处理完成的问题"),
        ]

    def _build_actions(self, issues: list[dict[str, str]]) -> list[str]:
        actions: list[str] = []
        for item in issues:
            suggestion = item["suggestion"].strip()
            if suggestion and suggestion not in actions:
                actions.append(suggestion)

        if any(item["level"] == "P0" for item in issues):
            actions.append("存在 P0 风险，建议暂停签署并升级到法务负责人或经营负责人审批。")

        return actions[:5]

    def _overall_risk(self, issues: list[dict[str, str]]) -> str:
        if not issues:
            return "P3"
        return min((item["level"] for item in issues), key=lambda item: SEVERITY_RANK.get(item, 99))

    def _normalize_text(self, value: str) -> str:
        return re.sub(r"\s+", "", value)

    def _split_clauses(self, text: str) -> list[str]:
        segments = re.split(r"[\n。；;]", text)
        return [item.strip() for item in segments if item.strip()]

    def _split_structured_clauses(self, text: str) -> list[dict[str, str]]:
        lines = [item.strip() for item in re.split(r"[\n\r]+", text) if item.strip()]
        if len(lines) <= 1:
            lines = [item.strip() for item in re.split(r"(?=第[一二三四五六七八九十0-9]+条)", text) if item.strip()]
        if len(lines) <= 1:
            lines = [item.strip() for item in re.split(r"[。；;]", text) if item.strip()]

        clauses: list[dict[str, str]] = []
        for index, line in enumerate(lines, start=1):
            title_match = re.match(r"^(第[一二三四五六七八九十0-9]+条)", line)
            title = title_match.group(1) if title_match else f"条款 {index}"
            content = line if line.endswith(("。", "；", ";")) else line
            clauses.append(
                {
                    "clause_no": index,
                    "title": title,
                    "content": content[:4000],
                    "source_ref": f"Clause {index}",
                }
            )
        return clauses

    def _extract_text_from_file(self, *, filename: str, file_bytes: bytes) -> str:
        suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if suffix in {"txt", "md"}:
            return file_bytes.decode("utf-8", errors="ignore").strip()

        if suffix == "pdf":
            reader = PdfReader(BytesIO(file_bytes))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
            return "\n".join([item for item in pages if item]).strip()

        if suffix == "docx":
            document = Document(BytesIO(file_bytes))
            paragraphs = [item.text.strip() for item in document.paragraphs if item.text.strip()]
            return "\n".join(paragraphs).strip()

        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unsupported file type. Please upload txt, pdf, or docx.",
        )

    def _matches_rule(self, rule: RuleDefinition, normalized_text: str) -> bool:
        normalized_patterns = [pattern.replace(" ", "") for pattern in rule.patterns]
        if rule.match_mode == "all":
            return all(pattern in normalized_text for pattern in normalized_patterns)
        return any(pattern in normalized_text for pattern in normalized_patterns)

    def _find_evidence(self, clauses: list[str], patterns: tuple[str, ...]) -> str:
        for clause in clauses:
            normalized_clause = self._normalize_text(clause)
            if any(pattern.replace(" ", "") in normalized_clause for pattern in patterns):
                return clause[:280]
        return ""

    def _build_semantic_issue(
        self,
        *,
        title: str,
        issue_type: str,
        level: str,
        document: str,
        evidence: str,
        detail: str,
        suggestion: str,
        rule_name: str,
    ) -> dict[str, str]:
        return {
            "title": title,
            "type": issue_type,
            "level": level,
            "status": "待处理",
            "document": document,
            "detail": detail,
            "evidence": evidence[:280],
            "suggestion": suggestion,
            "rule_name": rule_name,
        }

    def _merge_issues(self, issues: list[dict[str, str]]) -> list[dict[str, str]]:
        deduped: dict[tuple[str, str], dict[str, str]] = {}
        for item in issues:
            key = (item["title"], item["evidence"])
            deduped[key] = item

        merged = list(deduped.values())
        merged = self._dedupe_by_semantic_similarity(merged)
        return sorted(merged, key=self._sort_issue_dict)

    def _dedupe_by_semantic_similarity(self, issues: list[dict[str, str]]) -> list[dict[str, str]]:
        if len(issues) <= 1:
            return issues

        result: list[dict[str, str]] = []
        for issue in issues:
            if self._is_duplicate_of_existing(issue, result):
                continue
            result.append(issue)
        return result

    def _is_duplicate_of_existing(self, issue: dict[str, str], existing: list[dict[str, str]]) -> bool:
        issue_title = self._normalize_text(issue["title"])
        issue_evidence = self._normalize_text(issue.get("evidence", ""))
        issue_type = issue.get("type", "")

        for existing_issue in existing:
            existing_title = self._normalize_text(existing_issue["title"])
            existing_evidence = self._normalize_text(existing_issue.get("evidence", ""))
            existing_type = existing_issue.get("type", "")

            if self._titles_are_similar(issue_title, existing_title):
                if issue_type == existing_type:
                    return True
                if self._evidence_overlap(issue_evidence, existing_evidence):
                    return True

            if issue_evidence and existing_evidence and self._evidence_overlap(issue_evidence, existing_evidence):
                if issue_type == existing_type:
                    return True

        return False

    def _titles_are_similar(self, title1: str, title2: str) -> bool:
        if title1 == title2:
            return True

        keywords1 = set(self._extract_keywords(title1))
        keywords2 = set(self._extract_keywords(title2))
        if not keywords1 or not keywords2:
            return False

        intersection = keywords1 & keywords2
        smaller = min(len(keywords1), len(keywords2))
        return len(intersection) >= smaller * 0.6

    def _extract_keywords(self, text: str) -> list[str]:
        stop_words = {"的", "存在", "条款", "风险", "表述", "需", "法务", "确认", "可能", "会", "带来"}
        words = re.findall(r"[\u4e00-\u9fa5]{2,}", text)
        return [w for w in words if w not in stop_words]

    def _evidence_overlap(self, evidence1: str, evidence2: str) -> bool:
        if not evidence1 or not evidence2:
            return False

        min_len = min(len(evidence1), len(evidence2))
        if min_len < 10:
            return evidence1 == evidence2

        shorter = evidence1 if len(evidence1) <= len(evidence2) else evidence2
        longer = evidence2 if len(evidence1) <= len(evidence2) else evidence1
        return shorter in longer

    def _has_any(self, normalized_text: str, phrases: tuple[str, ...]) -> bool:
        return any(phrase.replace(" ", "") in normalized_text for phrase in phrases)

    def _infer_document_label(self, clause: str) -> str:
        normalized = self._normalize_text(clause)

        if self._has_any(normalized, ("付款", "结算", "回款", "预付款", "尾款", "发票")):
            return "付款条款"
        if self._has_any(normalized, ("验收", "交付", "上线", "整改")):
            return "交付与验收条款"
        if self._has_any(normalized, ("服务", "驻场", "响应", "运维", "支持")):
            return "服务条款"
        if self._has_any(normalized, ("变更", "调整", "新增需求", "范围")):
            return "变更条款"
        if self._has_any(normalized, ("责任", "赔偿", "违约", "损失")):
            return "责任条款"
        return "合同语义审查"

    def _issue_origin(self, rule_name: str) -> str:
        if rule_name.startswith("llm_"):
            return "LLM 语义审查"
        if rule_name.startswith("semantic_"):
            return "启发式语义审查"
        if rule_name == "no_rule_hit":
            return "规则摘要"
        return "规则命中"

    def _get_job_record(self, db: Session, job_id: str) -> ReviewJob:
        job = db.get(ReviewJob, job_id)
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Review job '{job_id}' not found.",
            )
        return job

    def _latest_job(self, db: Session) -> ReviewJob | None:
        return db.scalars(select(ReviewJob).order_by(ReviewJob.created_at.desc()).limit(1)).first()

    def _refresh_job_summary(self, db: Session, job_id: str) -> None:
        job = self._get_job_record(db, job_id)
        issues = db.scalars(select(ReviewIssueRecord).where(ReviewIssueRecord.job_id == job_id)).all()
        issue_dicts = [
            {
                "level": item.level,
                "status": item.status,
                "suggestion": item.suggestion,
            }
            for item in issues
        ]
        job.summary = [item.model_dump() for item in self._build_summary(issue_dicts)]
        job.review_actions = self._build_actions(issue_dicts)
        job.issue_count = len(issue_dicts)
        job.high_risk_issue_count = len([item for item in issue_dicts if item["level"] in {"P0", "P1"}])
        job.overall_risk = self._overall_risk(issue_dicts)
        job.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(job)

    def _to_job_response(self, job: ReviewJob) -> ReviewJobResponse:
        return ReviewJobResponse(
            id=job.id,
            project_id=job.project_id,
            contract_name=job.contract_name,
            contract_type=job.contract_type,
            trigger=job.trigger,
            status=job.status,
            overall_risk=job.overall_risk,
            issue_count=job.issue_count,
            high_risk_issue_count=job.high_risk_issue_count,
            summary=[MetricItem(**item) for item in job.summary],
            review_actions=list(job.review_actions),
            created_at=job.created_at,
            updated_at=job.updated_at,
            completed_at=job.completed_at,
        )

    def _sort_issue_dict(self, item: dict[str, str]) -> tuple[int, str]:
        return (SEVERITY_RANK.get(item["level"], 99), item["title"])

    def _sort_issue_record(self, item: ReviewIssueRecord) -> tuple[int, str]:
        return (SEVERITY_RANK.get(item.level, 99), item.title)


review_service = ReviewService()
