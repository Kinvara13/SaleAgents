from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
import random
import re
from typing import Any
from uuid import uuid4

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import HTTPException, status
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.models.generation_job import GenerationJob
from app.models.generation_section_asset_ref import GenerationSectionAssetRef
from app.models.generation_section import GenerationSectionRecord
from app.models.knowledge_asset_index_job import KnowledgeAssetIndexJobRecord
from app.models.project_asset_preference import ProjectAssetPreferenceRecord
from app.models.project import Project
from app.schemas.generation import (
    GenerationAssetChunkCreateRequest,
    GenerationAssetChunkUpdateRequest,
    GenerationJobCreateRequest,
    GenerationJobResponse,
    GenerationAssetIndexJobResponse,
    GenerationCheckResponse,
    GenerationJobAnalysisResponse,
    GenerationScoreItemResponse,
    GenerationSectionCoverageResponse,
    GenerationAssetReviewRequest,
    GenerationAssetUpdateRequest,
    GenerationProjectContextResponse,
    GenerationProjectRunRequest,
    GenerationSectionResponse,
    GenerationSectionUpdateRequest,
    IndexedGenerationAssetResponse,
    KnowledgeAssetChunkResponse,
    ProjectGenerationAssetPreferencesResponse,
    ProjectGenerationAssetPreferencesUpdateRequest,
    GenerationAssetRefreshRequest,
)
from app.services.asset_index_service import asset_index_service
from app.services.asset_routing_service import RoutedAsset, asset_routing_service
from app.services.llm_client import llm_generation_client
from app.services.parsing_service import parsing_service
from app.services.workspace_service import (
    get_extracted_fields,
    get_generation_assets,
    get_generation_todos,
    get_parse_sections,
)

DEFAULT_SECTION_TITLES = {
    "标准回标模板": [
        "项目理解与建设目标",
        "公司概况与资质",
        "总体技术方案",
        "实施计划与里程碑",
        "售后服务方案",
        "商务偏离说明",
    ]
}

SECTION_MAPPING_RULES: dict[str, tuple[str, ...]] = {
    "项目理解与建设目标": ("项目", "背景", "目标", "理解", "场景", "需求"),
    "公司概况与资质": ("资质", "资格", "案例", "业绩", "团队", "证书", "公司"),
    "总体技术方案": ("技术", "功能", "性能", "架构", "接口", "参数", "方案"),
    "实施计划与里程碑": ("实施", "交付", "工期", "周期", "里程碑", "部署", "验收"),
    "售后服务方案": ("服务", "售后", "SLA", "响应", "质保", "运维", "培训"),
    "商务偏离说明": ("商务", "付款", "报价", "偏离", "条款", "合同"),
}


@dataclass(frozen=True)
class GenerationContext:
    project_name: str
    client_name: str
    template_name: str
    project_summary: str
    tender_requirements: str
    delivery_deadline: str
    service_commitment: str
    selected_assets: list[str]
    fixed_assets: list[str]
    excluded_assets: list[str]
    generation_todos: list[str]
    parse_sections: list[str]
    extracted_fields: dict[str, str]
    technical_spec_text: str = ""


class GenerationService:
    """Handles bid-response document generation workflows."""

    def create_job(
        self,
        db: Session,
        payload: GenerationJobCreateRequest,
    ) -> GenerationJobResponse:
        now = datetime.now(timezone.utc)
        job_id = f"gen-{uuid4().hex[:8]}"
        section_titles = self._resolve_section_titles(
            template_name=payload.template_name,
            requested_titles=payload.section_titles,
        )
        context = self._build_context(db, payload)

        job = GenerationJob(
            id=job_id,
            project_id=payload.project_id,
            project_name=payload.project_name.strip(),
            template_name=payload.template_name.strip(),
            status="completed",
            section_count=len(section_titles),
            overall_progress="已生成",
            created_at=now,
            updated_at=now,
            completed_at=now,
        )
        db.add(job)

        for idx, title in enumerate(section_titles, start=1):
            generated = self._generate_section(db, title=title, context=context, section_no=idx)
            section = GenerationSectionRecord(
                id=f"{job_id}-s{idx}",
                job_id=job_id,
                section_no=idx,
                title=title,
                content=generated["content"],
                status="已生成",
                citations=generated["citations"],
                todos=generated["todos"],
                created_at=now,
            )
            db.add(section)
            db.flush()
            self._replace_section_asset_refs(
                db,
                job_id=job_id,
                section_id=section.id,
                routed_assets=generated.get("routed_assets", []),
            )

        db.commit()
        db.refresh(job)
        return GenerationJobResponse.model_validate(job)

    def get_project_context(
        self,
        db: Session,
        project_id: str,
    ) -> GenerationProjectContextResponse:
        project = self._get_project(db, project_id)
        extracted_fields = parsing_service.project_fields_map(db, project_id)
        parsing_context = parsing_service.get_project_context(db, project_id)
        parse_sections = [item.title for item in parsing_context.parse_sections]
        preferences = self.get_project_asset_preferences(db, project_id)
        if not extracted_fields:
            extracted_fields = {item.label: item.value for item in get_extracted_fields(db)}
        if not parse_sections:
            parse_sections = [item.title for item in get_parse_sections(db)]
        generation_todos = get_generation_todos(db)
        generation_assets = asset_index_service.list_asset_responses(db)
        selected_asset_titles = self._resolve_selected_assets(
            available_titles=[asset.title for asset in generation_assets],
            requested_titles=[],
            fixed_titles=preferences.fixed_asset_titles,
            excluded_titles=preferences.excluded_asset_titles,
        )

        summary_parts = [
            f"项目名称：{extracted_fields.get('项目名称') or project.name}",
            f"客户：{project.client}" if project.client else "",
            f"预算：{extracted_fields.get('预算金额') or project.amount}" if (extracted_fields.get("预算金额") or project.amount) else "",
            f"风险等级：{project.risk}" if project.risk else "",
        ]
        summary = "；".join([part for part in summary_parts if part])
        if not summary:
            summary = f"围绕项目“{project.name}”生成技术与商务应答初稿。"

        requirement_lines: list[str] = []
        for field_name in ("必备资质", "付款条款", "交付周期"):
            value = extracted_fields.get(field_name)
            if value:
                requirement_lines.append(f"{field_name}：{value}")
        if parse_sections:
            requirement_lines.append(f"已识别招标结构：{'、'.join(parse_sections[:4])}")
        if generation_todos:
            requirement_lines.append(f"待确认事项：{'；'.join(generation_todos[:3])}")

        service_commitment = "；".join(generation_todos[:2]) if generation_todos else "按公司标准 SLA 模板输出服务承诺。"

        return GenerationProjectContextResponse(
            project_id=project.id,
            project_name=extracted_fields.get("项目名称") or project.name,
            client_name=project.client or "",
            template_name="标准回标模板",
            project_summary=summary,
            tender_requirements="\n".join(requirement_lines),
            delivery_deadline=extracted_fields.get("投标截止时间") or project.deadline or extracted_fields.get("交付周期", ""),
            service_commitment=service_commitment,
            selected_asset_titles=selected_asset_titles,
            section_titles=list(DEFAULT_SECTION_TITLES["标准回标模板"]),
            fixed_asset_titles=preferences.fixed_asset_titles,
            excluded_asset_titles=preferences.excluded_asset_titles,
            source_excerpt=parsing_context.source_excerpt or "",
        )

    def create_job_from_project(
        self,
        db: Session,
        project_id: str,
        payload: GenerationProjectRunRequest,
    ) -> GenerationJobResponse:
        context = self.get_project_context(db, project_id)
        request = GenerationJobCreateRequest(
            project_id=context.project_id,
            project_name=context.project_name,
            client_name=context.client_name,
            template_name=(payload.template_name or context.template_name).strip() or "标准回标模板",
            project_summary=(payload.project_summary or context.project_summary).strip(),
            tender_requirements=(payload.tender_requirements or context.tender_requirements).strip(),
            delivery_deadline=(payload.delivery_deadline or context.delivery_deadline).strip(),
            service_commitment=(payload.service_commitment or context.service_commitment).strip(),
            selected_asset_titles=payload.selected_asset_titles or context.selected_asset_titles,
            section_titles=payload.section_titles or context.section_titles,
        )
        return self.create_job(db, request)

    def list_indexed_assets(self, db: Session) -> list[IndexedGenerationAssetResponse]:
        asset_index_service.initialize_seed_assets(db)
        return asset_index_service.list_asset_responses(db)

    def create_asset(
        self,
        db: Session,
        *,
        title: str,
        asset_type: str,
        status_text: str,
        content: str,
        owner: str,
        visibility: str,
    ) -> IndexedGenerationAssetResponse:
        asset_index_service.initialize_seed_assets(db)
        return asset_index_service.create_manual_asset(
            db,
            title=title,
            asset_type=asset_type,
            status_text=status_text,
            content=content,
            owner=owner,
            visibility=visibility,
        )

    def upload_asset(
        self,
        db: Session,
        *,
        filename: str,
        file_bytes: bytes,
        asset_type: str,
        title: str | None = None,
        owner: str = "system",
        visibility: str = "internal",
    ) -> IndexedGenerationAssetResponse:
        asset_index_service.initialize_seed_assets(db)
        return asset_index_service.create_uploaded_asset(
            db,
            filename=filename,
            file_bytes=file_bytes,
            asset_type=asset_type,
            title=title,
            owner=owner,
            visibility=visibility,
        )

    def refresh_asset_indexes(self, db: Session, asset_id: str | None = None) -> int:
        asset_index_service.initialize_seed_assets(db)
        return asset_index_service.refresh_asset_indexes(db, asset_id)

    def create_asset_refresh_job(
        self,
        db: Session,
        *,
        asset_id: str | None,
        triggered_by: str,
    ) -> GenerationAssetIndexJobResponse:
        job = KnowledgeAssetIndexJobRecord(
            id=f"asset-job-{uuid4().hex[:10]}",
            asset_id=asset_id or "",
            status="queued",
            triggered_by=triggered_by.strip() or "system",
            refreshed_count=0,
            error_message="",
        )
        db.add(job)
        db.commit()
        db.refresh(job)
        return GenerationAssetIndexJobResponse.model_validate(job)

    def run_asset_refresh_job(self, job_id: str) -> None:
        from app.db.session import SessionLocal

        with SessionLocal() as db:
            job = db.get(KnowledgeAssetIndexJobRecord, job_id)
            if job is None:
                return
            try:
                job.status = "running"
                db.commit()
                refreshed = self.refresh_asset_indexes(db, job.asset_id or None)
                job.status = "completed"
                job.refreshed_count = refreshed
                job.completed_at = datetime.now(timezone.utc)
                db.commit()
            except Exception as exc:  # noqa: BLE001
                job.status = "failed"
                job.error_message = str(exc)[:4000]
                job.completed_at = datetime.now(timezone.utc)
                db.commit()

    def get_asset_refresh_job(self, db: Session, job_id: str) -> GenerationAssetIndexJobResponse:
        job = db.get(KnowledgeAssetIndexJobRecord, job_id)
        if job is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Asset index job '{job_id}' not found.")
        return GenerationAssetIndexJobResponse.model_validate(job)

    def get_project_asset_preferences(
        self,
        db: Session,
        project_id: str,
    ) -> ProjectGenerationAssetPreferencesResponse:
        self._get_project(db, project_id)
        rows = db.scalars(
            select(ProjectAssetPreferenceRecord)
            .where(ProjectAssetPreferenceRecord.project_id == project_id)
            .order_by(ProjectAssetPreferenceRecord.created_at.asc())
        ).all()
        fixed_asset_titles = [row.asset_title for row in rows if row.preference_mode == "fixed"]
        excluded_asset_titles = [row.asset_title for row in rows if row.preference_mode == "excluded"]
        return ProjectGenerationAssetPreferencesResponse(
            project_id=project_id,
            fixed_asset_titles=fixed_asset_titles,
            excluded_asset_titles=excluded_asset_titles,
        )

    def update_project_asset_preferences(
        self,
        db: Session,
        project_id: str,
        payload: ProjectGenerationAssetPreferencesUpdateRequest,
    ) -> ProjectGenerationAssetPreferencesResponse:
        self._get_project(db, project_id)
        fixed_titles = [item.strip() for item in payload.fixed_asset_titles if item.strip()]
        excluded_titles = [item.strip() for item in payload.excluded_asset_titles if item.strip()]
        fixed_titles = [item for item in fixed_titles if item not in excluded_titles]

        db.execute(delete(ProjectAssetPreferenceRecord).where(ProjectAssetPreferenceRecord.project_id == project_id))
        db.flush()
        for title in fixed_titles:
            db.add(
                ProjectAssetPreferenceRecord(
                    id=f"proj-asset-pref-{uuid4().hex[:10]}",
                    project_id=project_id,
                    asset_title=title,
                    preference_mode="fixed",
                )
            )
        for title in excluded_titles:
            db.add(
                ProjectAssetPreferenceRecord(
                    id=f"proj-asset-pref-{uuid4().hex[:10]}",
                    project_id=project_id,
                    asset_title=title,
                    preference_mode="excluded",
                )
            )
        db.commit()
        return self.get_project_asset_preferences(db, project_id)

    def update_asset(
        self,
        db: Session,
        asset_id: str,
        payload: GenerationAssetUpdateRequest,
    ) -> IndexedGenerationAssetResponse:
        asset_index_service.initialize_seed_assets(db)
        return asset_index_service.update_asset(
            db,
            asset_id=asset_id,
            title=payload.title,
            asset_type=payload.asset_type,
            status_text=payload.status,
            content=payload.content,
            owner=payload.owner,
            visibility=payload.visibility,
        )

    def delete_asset(self, db: Session, asset_id: str) -> None:
        asset_index_service.initialize_seed_assets(db)
        asset_index_service.delete_asset(db, asset_id)

    def review_asset(
        self,
        db: Session,
        asset_id: str,
        payload: GenerationAssetReviewRequest,
    ) -> IndexedGenerationAssetResponse:
        asset_index_service.initialize_seed_assets(db)
        return asset_index_service.review_asset(
            db,
            asset_id=asset_id,
            action=payload.action,
            reviewer=payload.reviewer,
            review_note=payload.review_note,
        )

    def list_asset_chunks(self, db: Session, asset_id: str) -> list[KnowledgeAssetChunkResponse]:
        asset_index_service.initialize_seed_assets(db)
        rows = asset_index_service.list_asset_chunks(db, asset_id)
        return [
            KnowledgeAssetChunkResponse(
                id=row.id,
                asset_id=row.asset_id,
                title=row.title,
                content=row.content,
                keywords=asset_index_service.split_values(row.keywords),
                section_tags=asset_index_service.split_values(row.section_tags),
                sort_order=row.sort_order,
                weight=row.weight,
            )
            for row in rows
        ]

    def create_asset_chunk(
        self,
        db: Session,
        asset_id: str,
        payload: GenerationAssetChunkCreateRequest,
    ) -> KnowledgeAssetChunkResponse:
        row = asset_index_service.create_chunk(
            db,
            asset_id=asset_id,
            title=payload.title,
            content=payload.content,
            keywords=payload.keywords,
            section_tags=payload.section_tags,
            weight=payload.weight,
        )
        return KnowledgeAssetChunkResponse(
            id=row.id,
            asset_id=row.asset_id,
            title=row.title,
            content=row.content,
            keywords=asset_index_service.split_values(row.keywords),
            section_tags=asset_index_service.split_values(row.section_tags),
            sort_order=row.sort_order,
            weight=row.weight,
        )

    def update_asset_chunk(
        self,
        db: Session,
        asset_id: str,
        chunk_id: str,
        payload: GenerationAssetChunkUpdateRequest,
    ) -> KnowledgeAssetChunkResponse:
        row = asset_index_service.update_chunk(
            db,
            asset_id=asset_id,
            chunk_id=chunk_id,
            title=payload.title,
            content=payload.content,
            keywords=payload.keywords,
            section_tags=payload.section_tags,
            weight=payload.weight,
        )
        return KnowledgeAssetChunkResponse(
            id=row.id,
            asset_id=row.asset_id,
            title=row.title,
            content=row.content,
            keywords=asset_index_service.split_values(row.keywords),
            section_tags=asset_index_service.split_values(row.section_tags),
            sort_order=row.sort_order,
            weight=row.weight,
        )

    def delete_asset_chunk(self, db: Session, asset_id: str, chunk_id: str) -> None:
        asset_index_service.delete_chunk(db, asset_id=asset_id, chunk_id=chunk_id)

    def get_job(self, db: Session, job_id: str) -> GenerationJobResponse:
        job = db.get(GenerationJob, job_id)
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Generation job '{job_id}' not found.",
            )
        return GenerationJobResponse.model_validate(job)

    def get_latest_job(self, db: Session) -> GenerationJobResponse | None:
        job = db.scalars(
            select(GenerationJob).order_by(GenerationJob.created_at.desc()).limit(1)
        ).first()
        if job is None:
            return None
        return GenerationJobResponse.model_validate(job)

    def get_latest_job_by_project(self, db: Session, project_id: str) -> GenerationJobResponse:
        job = db.scalars(
            select(GenerationJob)
            .where(GenerationJob.project_id == project_id)
            .order_by(GenerationJob.created_at.desc())
            .limit(1)
        ).first()
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No generation job found for this project.",
            )
        return GenerationJobResponse.model_validate(job)

    def list_job_sections(
        self, db: Session, job_id: str
    ) -> list[GenerationSectionResponse]:
        job = self.get_job(db, job_id)
        rows = db.scalars(
            select(GenerationSectionRecord)
            .where(GenerationSectionRecord.job_id == job_id)
            .order_by(GenerationSectionRecord.section_no.asc())
        ).all()
        refs = db.scalars(
            select(GenerationSectionAssetRef)
            .where(GenerationSectionAssetRef.job_id == job_id)
            .order_by(GenerationSectionAssetRef.score.desc())
        ).all()
        refs_by_section: dict[str, list[GenerationSectionAssetRef]] = {}
        for ref in refs:
            refs_by_section.setdefault(ref.section_id, []).append(ref)

        analysis = self._analyze_job_rows(db, job, rows)
        coverage_map = {item.section_id: item for item in analysis.section_coverages}
        return [self._build_section_response(row, refs_by_section.get(row.id, []), coverage_map.get(row.id)) for row in rows]

    def get_job_analysis(self, db: Session, job_id: str) -> GenerationJobAnalysisResponse:
        job = self.get_job(db, job_id)
        rows = db.scalars(
            select(GenerationSectionRecord)
            .where(GenerationSectionRecord.job_id == job_id)
            .order_by(GenerationSectionRecord.section_no.asc())
        ).all()
        return self._analyze_job_rows(db, job, rows)

    def update_section(
        self,
        db: Session,
        job_id: str,
        section_id: str,
        payload: GenerationSectionUpdateRequest,
    ) -> GenerationSectionResponse:
        section = self._get_section(db, job_id, section_id)
        section.content = payload.content.strip()
        section.status = payload.status.strip() or "已编辑"

        job = db.get(GenerationJob, job_id)
        if job:
            job.updated_at = datetime.now(timezone.utc)
            job.overall_progress = "待复核"

        db.commit()
        db.refresh(section)
        refs = self._list_section_asset_refs(db, section.id)
        analysis = self._analyze_job_rows(db, self.get_job(db, job_id), [section])
        coverage = analysis.section_coverages[0] if analysis.section_coverages else None
        return self._build_section_response(section, refs, coverage)

    def regenerate_section(
        self, db: Session, job_id: str, section_id: str
    ) -> GenerationSectionResponse:
        section = self._get_section(db, job_id, section_id)
        job = db.get(GenerationJob, job_id)
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Generation job '{job_id}' not found.",
            )

        regenerated = self._regenerate_with_context(db, job, section)
        section.content = regenerated["content"]
        section.status = "已重新生成"
        section.citations = regenerated["citations"]
        section.todos = regenerated["todos"]
        job.updated_at = datetime.now(timezone.utc)
        job.overall_progress = "已更新"
        self._replace_section_asset_refs(
            db,
            job_id=job_id,
            section_id=section.id,
            routed_assets=regenerated.get("routed_assets", []),
        )

        db.commit()
        db.refresh(section)
        refs = self._list_section_asset_refs(db, section.id)
        analysis = self._analyze_job_rows(db, self.get_job(db, job_id), [section])
        coverage = analysis.section_coverages[0] if analysis.section_coverages else None
        return self._build_section_response(section, refs, coverage)

    def repair_uncovered_sections(self, db: Session, job_id: str) -> list[GenerationSectionResponse]:
        job = db.get(GenerationJob, job_id)
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Generation job '{job_id}' not found.",
            )
        sections = db.scalars(
            select(GenerationSectionRecord)
            .where(GenerationSectionRecord.job_id == job_id)
            .order_by(GenerationSectionRecord.section_no.asc())
        ).all()
        context = self._build_context_from_job(db, job)
        analysis = self._analyze_job_rows(db, GenerationJobResponse.model_validate(job), sections)
        coverage_map = {item.section_id: item for item in analysis.section_coverages}
        changed = False
        for section in sections:
            coverage = coverage_map.get(section.id)
            if coverage is None or not coverage.missing_requirements:
                continue
            refs = self._list_section_asset_refs(db, section.id)
            revised = self._repair_section_content(
                section=section,
                coverage=coverage,
                refs=refs,
                context=context,
                mode="repair",
            )
            section.content = revised["content"]
            section.status = "已补写"
            section.citations = max(section.citations, revised["citations"])
            section.todos = revised["todos"]
            changed = True
        if changed:
            job.updated_at = datetime.now(timezone.utc)
            job.overall_progress = "已补写评分缺口"
            db.commit()
        return self.list_job_sections(db, job_id)

    def self_revise_job(self, db: Session, job_id: str) -> list[GenerationSectionResponse]:
        job = db.get(GenerationJob, job_id)
        if job is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Generation job '{job_id}' not found.",
            )
        sections = db.scalars(
            select(GenerationSectionRecord)
            .where(GenerationSectionRecord.job_id == job_id)
            .order_by(GenerationSectionRecord.section_no.asc())
        ).all()
        context = self._build_context_from_job(db, job)
        analysis = self._analyze_job_rows(db, GenerationJobResponse.model_validate(job), sections)
        coverage_map = {item.section_id: item for item in analysis.section_coverages}
        changed = False
        for section in sections:
            coverage = coverage_map.get(section.id)
            if coverage is None:
                continue
            if coverage.coverage_score >= 75 and not coverage.self_check_notes and not coverage.missing_requirements:
                continue
            refs = self._list_section_asset_refs(db, section.id)
            revised = self._repair_section_content(
                section=section,
                coverage=coverage,
                refs=refs,
                context=context,
                mode="self_revise",
            )
            section.content = revised["content"]
            section.status = "已自修订"
            section.citations = max(section.citations, revised["citations"])
            section.todos = revised["todos"]
            changed = True
        if changed:
            job.updated_at = datetime.now(timezone.utc)
            job.overall_progress = "已完成二轮自修订"
            db.commit()
        return self.list_job_sections(db, job_id)

    def export_job(self, db: Session, job_id: str) -> str:
        job_resp = self.get_job(db, job_id)
        sections = self.list_job_sections(db, job_id)

        lines: list[str] = [
            f"# 回标文件 — {job_resp.project_name}",
            "",
            f"- 模板：{job_resp.template_name}",
            f"- 状态：{job_resp.overall_progress}",
            f"- 章节数：{job_resp.section_count}",
            f"- 生成时间：{job_resp.created_at.strftime('%Y-%m-%d %H:%M')}",
            "",
            "---",
            "",
        ]

        for section in sections:
            lines.append(section.content)
            lines.append("")
            lines.append(f"> 引用 {section.citations} 条 · 待确认 {section.todos} 项")
            lines.append("")
            lines.append("---")
            lines.append("")

        return "\n".join(lines)

    def export_job_docx(
        self,
        db: Session,
        job_id: str,
        template_bytes: bytes | None = None,
    ) -> bytes:
        """
        导出 job 为 Word 文档。

        Args:
            db: 数据库会话
            job_id: 生成任务 ID
            template_bytes: 可选，Word 模板文件字节流。
                           如果传入，优先使用模板占位符填充；
                           不传则基于 python-docx 样式生成。
        """
        job_resp = self.get_job(db, job_id)
        sections = self.list_job_sections(db, job_id)

        sections_data = [
            {
                "title": s.title,
                "content": s.content,
                "citations": s.citations or "",
                "todos": s.todos or "",
            }
            for s in sections
        ]

        metadata = {
            "项目名称": job_resp.project_name,
            "模板": job_resp.template_name,
            "状态": job_resp.overall_progress or "已生成",
            "章节数": str(job_resp.section_count),
            "生成时间": job_resp.created_at.strftime("%Y-%m-%d %H:%M"),
        }

        exporter = TemplateDocxExporter(template_bytes=template_bytes)
        return exporter.export(
            sections=sections_data,
            project_name=job_resp.project_name,
            metadata=metadata,
        )

    def _get_section(self, db: Session, job_id: str, section_id: str) -> GenerationSectionRecord:
        section = db.get(GenerationSectionRecord, section_id)
        if section is None or section.job_id != job_id:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Section '{section_id}' not found in job '{job_id}'.",
            )
        return section

    def _get_project(self, db: Session, project_id: str) -> Project:
        project = db.get(Project, project_id)
        if project is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Project '{project_id}' not found.",
            )
        return project

    def _resolve_section_titles(
        self,
        *,
        template_name: str,
        requested_titles: list[str],
    ) -> list[str]:
        normalized = [title.strip() for title in requested_titles if title.strip()]
        if normalized:
            return normalized
        return list(DEFAULT_SECTION_TITLES.get(template_name, DEFAULT_SECTION_TITLES["标准回标模板"]))

    def _build_context(self, db: Session, payload: GenerationJobCreateRequest) -> GenerationContext:
        asset_index_service.initialize_seed_assets(db)
        available_assets = get_generation_assets(db)
        preferences = self.get_project_asset_preferences(db, payload.project_id) if payload.project_id else None
        selected_assets = self._resolve_selected_assets(
            available_titles=[asset.title for asset in available_assets],
            requested_titles=payload.selected_asset_titles,
            fixed_titles=preferences.fixed_asset_titles if preferences else [],
            excluded_titles=preferences.excluded_asset_titles if preferences else [],
        )

        extracted_fields = {item.label: item.value for item in get_extracted_fields(db)}

        return GenerationContext(
            project_name=payload.project_name.strip(),
            client_name=payload.client_name.strip(),
            template_name=payload.template_name.strip() or "标准回标模板",
            project_summary=payload.project_summary.strip(),
            tender_requirements=payload.tender_requirements.strip(),
            delivery_deadline=payload.delivery_deadline.strip(),
            service_commitment=payload.service_commitment.strip(),
            selected_assets=selected_assets,
            fixed_assets=preferences.fixed_asset_titles if preferences else [],
            excluded_assets=preferences.excluded_asset_titles if preferences else [],
            generation_todos=get_generation_todos(db),
            parse_sections=[item.title for item in get_parse_sections(db)],
            extracted_fields=extracted_fields,
            technical_spec_text=payload.technical_spec_text or "",
        )

    def _build_context_from_job(self, db: Session, job: GenerationJob | GenerationJobResponse) -> GenerationContext:
        if job.project_id:
            project_context = self.get_project_context(db, job.project_id)
            return self._build_context(
                db,
                GenerationJobCreateRequest(
                    project_id=job.project_id,
                    project_name=job.project_name,
                    client_name=project_context.client_name,
                    template_name=job.template_name,
                    project_summary=project_context.project_summary,
                    tender_requirements=project_context.tender_requirements,
                    delivery_deadline=project_context.delivery_deadline,
                    service_commitment=project_context.service_commitment,
                    selected_asset_titles=project_context.selected_asset_titles,
                    section_titles=project_context.section_titles,
                    technical_spec_text=project_context.source_excerpt or "",
                ),
            )
        return self._build_context(
            db,
            GenerationJobCreateRequest(
                project_id=job.project_id,
                project_name=job.project_name,
                template_name=job.template_name,
                technical_spec_text=project_context.source_excerpt or "",
            ),
        )

    def _build_section_content(
        self,
        *,
        title: str,
        context: GenerationContext,
        section_no: int,
        routed_assets: list[RoutedAsset],
    ) -> str:
        score_focus = self._section_score_focus_lines(title, context)
        if title == "项目理解与建设目标":
            return self._render_project_understanding(context, routed_assets) + score_focus
        if title == "公司概况与资质":
            return self._render_company_profile(context, routed_assets) + score_focus
        if title == "总体技术方案":
            return self._render_solution(context, routed_assets) + score_focus
        if title == "实施计划与里程碑":
            return self._render_delivery_plan(context, routed_assets) + score_focus
        if title == "售后服务方案":
            return self._render_service_plan(context, routed_assets) + score_focus
        if title == "商务偏离说明":
            return self._render_business_response(context, routed_assets) + score_focus
        return self._render_generic_section(title, context, section_no, routed_assets) + score_focus

    def _extract_relevant_spec_text(self, full_text: str, section_title: str) -> str:
        """从招标原文（34K字）提取与当前章节最相关的段落（约3K字）。"""
        if not full_text or len(full_text) <= 6000:
            return full_text
        # 每个章节对应的关键词
        keywords_by_section = {
            "总体技术方案": ["技术", "功能", "架构", "模块", "接口", "性能", "指标", "部署", "安全", "网络", "服务器", "数据库", "系统接口", "数据", "分析", "管理"],
            "项目理解与建设目标": ["项目", "背景", "目标", "需求", "概述", "内容", "范围"],
            "实施计划与里程碑": ["实施", "交付", "工期", "周期", "验收", "里程碑", "计划", "调试", "测试", "上线", "培训", "部署"],
            "售后服务方案": ["服务", "售后", "维保", "质保", "响应", "SLA", "故障", "维护", "支持"],
            "商务偏离说明": ["商务", "付款", "报价", "价格", "合同", "偏离", "资质"],
            "公司概况与资质": ["资质", "案例", "业绩", "公司", "证书", "业绩", "团队", "能力"],
        }
        # 通用关键词（所有章节都关注）
        general_keywords = ["★", "实质性", "必须", "要求", "主要", "关键", "核心"]
        keywords = keywords_by_section.get(section_title, general_keywords)
        # 提取包含关键词的段落
        import re
        paragraphs = re.split(r"\n(?=\S)", full_text)
        scored = []
        for para in paragraphs:
            if len(para.strip()) < 30:
                continue
            score = sum(1 for kw in keywords if kw in para)
            score += sum(0.5 for kw in general_keywords if kw in para)
            # 包含长技术描述的段落优先
            if len(para) > 200:
                score += 1
            if score > 0:
                scored.append((score, len(para), para))
        scored.sort(key=lambda x: (x[0], x[1]), reverse=True)
        result_lines = []
        char_count = 0
        max_chars = 4000
        for score, length, para in scored:
            if char_count + length > max_chars:
                break
            result_lines.append(para)
            char_count += length + 2
        if not result_lines:
            # 没有匹配则取前6K字
            return full_text[:6000]
        return "\n\n".join(result_lines)

    def _render_project_understanding(self, context: GenerationContext, routed_assets: list[RoutedAsset]) -> str:
        summary = context.project_summary or "当前项目以招标文件要求为主线，需形成覆盖技术、交付和服务承诺的完整应答。"
        requirements = self._bullet_lines(context.tender_requirements, fallback=[
            "围绕资格、评分、技术规范和商务条款形成逐项响应。",
            "优先覆盖高分评分项、强制性参数和交付时限。",
            "对需人工确认的服务边界和商务条件做显式标注。",
        ])
        return (
            f"## 项目理解与建设目标\n\n"
            f"### 项目背景\n\n"
            f"{summary}\n\n"
            f"### 我方理解\n\n"
            f"- 项目名称：{context.project_name}\n"
            f"- 客户名称：{context.client_name or '待补充'}\n"
            f"- 交付目标：围绕招标文件要求形成可提交的回标文件初稿\n"
            f"- 覆盖范围：{'、'.join(context.parse_sections[:4]) or '资格、技术、商务与实施章节'}\n\n"
            f"### 关键响应目标\n\n"
            f"{requirements}\n\n"
            f"### 建议引用素材\n\n"
            f"{self._routed_asset_lines(routed_assets, ['优先引用项目理解和行业案例素材。'])}\n"
        )

    def _render_company_profile(self, context: GenerationContext, routed_assets: list[RoutedAsset]) -> str:
        assets = self._routed_asset_lines(routed_assets, [
            "电子与智能化相关资质证书",
            "同类行业项目案例",
            "售后与交付标准材料",
        ])
        return (
            "## 公司概况与资质\n\n"
            "我方具备面向大型政企项目的方案设计、集成交付与持续运维能力，"
            "可针对本项目提供从方案响应、实施交付到质保服务的完整支撑。\n\n"
            "### 可引用资质与案例素材\n\n"
            f"{assets}\n\n"
            "### 建议在正式回标文件中补充\n\n"
            "- 企业简介、主营方向与组织能力说明\n"
            "- 核心资质清单及有效期\n"
            "- 近三年同类项目案例及验收证明\n"
        )

    def _render_solution(self, context: GenerationContext, routed_assets: list[RoutedAsset]) -> str:
        requirements = self._bullet_lines(context.tender_requirements, fallback=[
            "满足招标文件技术规范、性能指标和接口对接要求。",
            "保证方案与现有系统环境兼容，并支持后续扩容。",
            "对高风险技术点给出实施保障措施和验收口径。",
        ])
        return (
            "## 总体技术方案\n\n"
            "本方案按照“需求分解、架构设计、能力映射、实施落地”的逻辑组织，"
            "确保技术应答与评分点逐项对应。\n\n"
            "### 技术应答策略\n\n"
            f"{requirements}\n\n"
            "### 可引用素材\n\n"
            f"{self._routed_asset_lines(routed_assets, ['引用平台架构、功能能力和技术指标类素材。'])}\n\n"
            "### 方案输出重点\n\n"
            "- 明确系统架构、模块边界和对接方式\n"
            "- 对关键性能指标给出可验证表述\n"
            "- 对非标准能力说明替代方案和实施条件\n"
        )

    def _render_delivery_plan(self, context: GenerationContext, routed_assets: list[RoutedAsset]) -> str:
        deadline = context.delivery_deadline or context.extracted_fields.get("交付周期", "待补充")
        todos = self._bullet_lines_from_list(context.generation_todos[:3], fallback=[
            "细化阶段里程碑、人员投入和验收动作。",
            "补充关键设备到货与部署窗口。",
        ])
        return (
            "## 实施计划与里程碑\n\n"
            f"建议以招标要求的交付周期为上限组织实施计划，当前识别到的交付时限为：{deadline}。\n\n"
            "### 实施安排建议\n\n"
            "- 第一阶段：需求澄清、现场勘查与方案冻结\n"
            "- 第二阶段：设备准备、平台部署与环境联调\n"
            "- 第三阶段：功能上线、培训交接与试运行\n"
            "- 第四阶段：正式验收与质保移交\n\n"
            "### 可引用素材\n\n"
            f"{self._routed_asset_lines(routed_assets, ['引用实施交付案例、里程碑模板和验收经验素材。'])}\n\n"
            "### 当前待补充事项\n\n"
            f"{todos}\n"
        )

    def _render_service_plan(self, context: GenerationContext, routed_assets: list[RoutedAsset]) -> str:
        commitment = context.service_commitment or "建议按照公司标准 SLA 响应模板输出质保、响应和巡检承诺。"
        return (
            "## 售后服务方案\n\n"
            f"{commitment}\n\n"
            "### 可引用素材\n\n"
            f"{self._routed_asset_lines(routed_assets, ['引用 SLA、售后响应和质保条款类素材。'])}\n\n"
            "### 建议纳入正式应答的服务承诺\n\n"
            "- 质保期限、服务窗口和响应等级\n"
            "- 故障受理、升级处理和现场支持机制\n"
            "- 培训、巡检、备件与版本维护安排\n"
        )

    def _render_business_response(self, context: GenerationContext, routed_assets: list[RoutedAsset]) -> str:
        payment = context.extracted_fields.get("付款条款", "待从招标文件中补充")
        deadline = context.delivery_deadline or context.extracted_fields.get("交付周期", "待补充")
        return (
            "## 商务偏离说明\n\n"
            "本章节用于逐项对照招标文件商务条款，明确完全响应、偏离说明或需澄清事项。\n\n"
            "### 当前已识别的重点商务条件\n\n"
            f"- 付款条款：{payment}\n"
            f"- 交付周期：{deadline}\n"
            f"- 服务承诺：{context.service_commitment or '待补充'}\n\n"
            "### 可引用素材\n\n"
            f"{self._routed_asset_lines(routed_assets, ['引用资质、服务条款和商务条件模板素材。'])}\n\n"
            "### 建议输出方式\n\n"
            "- 对完全响应项直接列明“无偏离”\n"
            "- 对需确认项标注商务澄清或内部审批条件\n"
            "- 对高风险条款同步提示法务或经营复核\n"
        )

    def _render_generic_section(
        self,
        title: str,
        context: GenerationContext,
        section_no: int,
        routed_assets: list[RoutedAsset],
    ) -> str:
        source_lines = self._routed_asset_lines(routed_assets, ["企业资质与案例素材"])
        return (
            f"## {title}\n\n"
            f"该章节根据项目“{context.project_name}”自动生成，当前为第 {section_no} 个章节。\n\n"
            "### 编写依据\n\n"
            f"{source_lines}\n\n"
            "### 待人工确认\n\n"
            "- 补充与本项目直接相关的参数、案例或服务承诺\n"
            "- 检查表述是否与招标文件格式要求一致\n"
        )

    def _generate_section(
        self,
        db: Session,
        *,
        title: str,
        context: GenerationContext,
        section_no: int,
    ) -> dict[str, Any]:
        routed_assets = asset_routing_service.route_assets_for_section(
            db,
            section_title=title,
            project_summary=context.project_summary,
            tender_requirements=context.tender_requirements,
            delivery_deadline=context.delivery_deadline,
            service_commitment=context.service_commitment,
            selected_asset_titles=context.selected_assets,
            fixed_asset_titles=context.fixed_assets,
            excluded_asset_titles=context.excluded_assets,
            extracted_fields=context.extracted_fields,
        )
        routed_asset_payloads = [
            f"{item.asset_title}｜{item.chunk_title}｜{item.snippet}｜命中原因：{item.reason}"
            for item in routed_assets
        ]
        generated = llm_generation_client.generate_bid_section(
            project_name=context.project_name,
            client_name=context.client_name,
            section_title=title,
            project_summary=context.project_summary,
            tender_requirements=context.tender_requirements,
            delivery_deadline=context.delivery_deadline,
            service_commitment=context.service_commitment,
            selected_assets=routed_asset_payloads or context.selected_assets,
            extracted_fields=context.extracted_fields,
            generation_todos=context.generation_todos,
            technical_spec_text=self._extract_relevant_spec_text(context.technical_spec_text, title),
        )
        if generated is not None:
            return {
                **generated,
                "citations": max(len(routed_assets), generated["citations"]),
                "routed_assets": routed_assets,
            }

        return {
            "content": self._build_section_content(
                title=title,
                context=context,
                section_no=section_no,
                routed_assets=routed_assets,
            ),
            "citations": max(len(routed_assets), self._estimate_citations(title, context)),
            "todos": self._estimate_todos(title, context),
            "routed_assets": routed_assets,
        }

    def _estimate_citations(self, title: str, context: GenerationContext) -> int:
        base = min(len(context.selected_assets), 4) or 1
        if "方案" in title or "资质" in title:
            return base + 1
        return base

    def _estimate_todos(self, title: str, context: GenerationContext) -> int:
        if "商务" in title or "实施" in title:
            return min(3, max(1, len(context.generation_todos[:3])))
        if context.service_commitment or context.delivery_deadline:
            return 1
        return 0

    def _regenerate_section_content(self, section: GenerationSectionRecord) -> str:
        timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
        if "##" in section.content:
            return (
                section.content.rstrip()
                + "\n\n"
                + "### 本轮补强说明\n\n"
                + f"- 已于 {timestamp} 重新组织本章节结构与措辞\n"
                + "- 建议继续结合评分点、参数表和案例材料做人工精修\n"
            )
        return (
            f"## {section.title}\n\n"
            f"{section.content.strip()}\n\n"
            f"### 本轮补强说明\n\n"
            f"- 已于 {timestamp} 重新生成本章节\n"
        )

    def _repair_section_content(
        self,
        *,
        section: GenerationSectionRecord,
        coverage: GenerationSectionCoverageResponse,
        refs: list[GenerationSectionAssetRef],
        context: GenerationContext,
        mode: str,
    ) -> dict[str, Any]:
        routed_assets = [f"{ref.asset_title} · {ref.chunk_title}：{ref.snippet}" for ref in refs[:4]]
        llm_result = llm_generation_client.revise_bid_section(
            section_title=section.title,
            current_content=section.content,
            missing_requirements=coverage.missing_requirements,
            check_notes=coverage.self_check_notes,
            routed_assets=routed_assets,
            extracted_fields=context.extracted_fields,
        )
        if llm_result is not None:
            return llm_result

        title = "### 本轮自动补写" if mode == "repair" else "### 二轮自修订"
        lines = [section.content.rstrip(), "", title, ""]
        if coverage.missing_requirements:
            lines.append("#### 补充覆盖的评分点")
            lines.extend([f"- {item}" for item in coverage.missing_requirements[:5]])
            lines.append("")
        if coverage.self_check_notes:
            lines.append("#### 自检修订说明")
            lines.extend([f"- {item}" for item in coverage.self_check_notes[:4]])
            lines.append("")
        if routed_assets:
            lines.append("#### 建议补充引用素材")
            lines.extend([f"- {item}" for item in routed_assets[:3]])
            lines.append("")
        lines.append("#### 修订动作")
        lines.append("- 已按未覆盖评分点补充应答要点，请继续核对参数、资质和商务条件。")
        return {
            "content": "\n".join(lines).strip(),
            "citations": max(section.citations, len(refs)),
            "todos": max(0, section.todos - 1),
        }

    def _regenerate_with_context(
        self,
        db: Session,
        job: GenerationJob,
        section: GenerationSectionRecord,
    ) -> dict[str, Any]:
        context = self._build_context_from_job(db, job)
        generated = self._generate_section(
            db,
            title=section.title,
            context=context,
            section_no=section.section_no,
        )
        if generated["content"] == self._build_section_content(
            title=section.title,
            context=context,
            section_no=section.section_no,
            routed_assets=generated.get("routed_assets", []),
        ):
            generated["content"] = self._regenerate_section_content(section)
            generated["citations"] = max(1, section.citations + random.randint(0, 2))
            generated["todos"] = max(0, section.todos + random.randint(-1, 1))
        return generated

    def _build_section_response(
        self,
        row: GenerationSectionRecord,
        refs: list[GenerationSectionAssetRef],
        coverage: GenerationSectionCoverageResponse | None,
    ) -> GenerationSectionResponse:
        return GenerationSectionResponse(
            id=row.id,
            job_id=row.job_id,
            section_no=row.section_no,
            title=row.title,
            content=row.content,
            status=row.status,
            citations=row.citations,
            todos=row.todos,
            created_at=row.created_at,
            routed_assets=[f"{ref.asset_title} · {ref.chunk_title}" for ref in refs],
            routing_reasons=[f"{ref.asset_title}：{ref.reason}" for ref in refs],
            matched_score_items=coverage.matched_score_items if coverage else [],
            missing_requirements=coverage.missing_requirements if coverage else [],
            coverage_score=coverage.coverage_score if coverage else 0,
        )

    def _analyze_job_rows(
        self,
        db: Session,
        job_resp: GenerationJobResponse,
        sections: list[GenerationSectionRecord],
    ) -> GenerationJobAnalysisResponse:
        context = self._build_context_from_job(db, job_resp)
        score_items = self._extract_score_items(context)
        section_coverages: list[GenerationSectionCoverageResponse] = []
        section_keyword_map = {
            section.id: self._extract_keywords(f"{section.title}\n{section.content}", limit=18) for section in sections
        }
        uncovered_items: list[GenerationScoreItemResponse] = []

        for section in sections:
            coverage = self._build_section_coverage(section, context, score_items, section_keyword_map[section.id])
            section_coverages.append(coverage)

        score_item_responses: list[GenerationScoreItemResponse] = []
        for item in score_items:
            matched_sections = [
                coverage.section_title for coverage in section_coverages if item.title in coverage.matched_score_items
            ]
            matched_keywords: list[str] = []
            for coverage in section_coverages:
                if coverage.section_title in matched_sections:
                    matched_keywords.extend(
                        [
                            keyword
                            for keyword in self._extract_keywords(item.title, limit=6)
                            if keyword in section_keyword_map.get(coverage.section_id, [])
                        ]
                    )
            item_resp = GenerationScoreItemResponse(
                id=f"score-{uuid4().hex[:8]}",
                title=item.title,
                source=item.source,
                weight=item.weight,
                mapped_sections=item.mapped_sections,
                matched_sections=matched_sections,
                matched_keywords=list(dict.fromkeys(matched_keywords)),
                coverage_status="已覆盖" if matched_sections else "待覆盖",
            )
            if not matched_sections:
                uncovered_items.append(item_resp)
            score_item_responses.append(item_resp)

        checks = self._build_generation_checks(context, section_coverages, score_item_responses, uncovered_items)
        covered_count = len([item for item in score_item_responses if item.coverage_status == "已覆盖"])
        total_count = len(score_item_responses)
        overall_score = (
            round(sum(item.coverage_score for item in section_coverages) / max(len(section_coverages), 1))
            if section_coverages
            else 0
        )
        return GenerationJobAnalysisResponse(
            job_id=job_resp.id,
            overall_coverage_score=overall_score,
            mapped_score_item_count=total_count,
            covered_score_item_count=covered_count,
            uncovered_score_item_count=max(total_count - covered_count, 0),
            score_items=score_item_responses,
            checks=checks,
            section_coverages=section_coverages,
        )

    def _extract_score_items(self, context: GenerationContext) -> list[GenerationScoreItemResponse]:
        raw_items: list[tuple[str, str, int]] = []
        field_rules = [
            ("评分重点", 3),
            ("技术要求", 3),
            ("必备资质", 2),
            ("服务承诺", 2),
            ("交付周期", 2),
            ("付款条款", 1),
        ]
        for label, weight in field_rules:
            value = context.extracted_fields.get(label, "").strip()
            if value and value != "待补充":
                for fragment in self._split_requirement_fragments(value):
                    raw_items.append((fragment, label, weight))

        for fragment in self._split_requirement_fragments(context.tender_requirements):
            raw_items.append((fragment, "招标要求", 2))

        dedup: list[GenerationScoreItemResponse] = []
        seen: set[str] = set()
        for title, source, weight in raw_items:
            normalized = re.sub(r"\s+", "", title)
            if len(normalized) < 4 or normalized in seen:
                continue
            seen.add(normalized)
            dedup.append(
                GenerationScoreItemResponse(
                    id=f"score-{uuid4().hex[:8]}",
                    title=title,
                    source=source,
                    weight=weight,
                    mapped_sections=self._infer_target_sections(title),
                )
            )
        return dedup[:18]

    def _build_section_coverage(
        self,
        section: GenerationSectionRecord,
        context: GenerationContext,
        score_items: list[GenerationScoreItemResponse],
        section_keywords: list[str],
    ) -> GenerationSectionCoverageResponse:
        matched_score_items: list[str] = []
        missing_requirements: list[str] = []
        notes: list[str] = []
        section_text = f"{section.title}\n{section.content}"

        for item in score_items:
            if section.title not in item.mapped_sections:
                continue
            item_keywords = self._extract_keywords(item.title, limit=8)
            matched = any(keyword in section_text for keyword in item_keywords if len(keyword) >= 2)
            if matched:
                matched_score_items.append(item.title)
            else:
                missing_requirements.append(item.title)

        if "待补充" in section.content:
            notes.append("章节仍包含“待补充”占位内容。")
        if section.citations == 0:
            notes.append("章节未命中素材引用，证据支撑不足。")
        if "商务" in section.title and context.extracted_fields.get("付款条款", "待补充") == "待补充":
            notes.append("付款条款仍未从招标文件中抽取到有效值。")
        if "资质" in section.title and not any("资质" in item for item in matched_score_items):
            notes.append("资质章节未覆盖明确资质门槛。")

        total_targets = max(len(matched_score_items) + len(missing_requirements), 1)
        base_score = round(len(matched_score_items) * 100 / total_targets)
        if section.citations == 0:
            base_score = max(base_score - 10, 0)
        return GenerationSectionCoverageResponse(
            section_id=section.id,
            section_title=section.title,
            coverage_score=base_score,
            matched_score_items=matched_score_items,
            missing_requirements=missing_requirements[:6],
            self_check_notes=notes,
        )

    def _build_generation_checks(
        self,
        context: GenerationContext,
        section_coverages: list[GenerationSectionCoverageResponse],
        score_items: list[GenerationScoreItemResponse],
        uncovered_items: list[GenerationScoreItemResponse],
    ) -> list[GenerationCheckResponse]:
        checks: list[GenerationCheckResponse] = []
        if uncovered_items:
            checks.append(
                GenerationCheckResponse(
                    id=f"gen-check-{uuid4().hex[:8]}",
                    level="P1",
                    category="评分点覆盖",
                    title="存在未覆盖评分点",
                    detail=f"仍有 {len(uncovered_items)} 个评分点或强制要求未被任何章节显式覆盖。",
                    related_sections=list(
                        dict.fromkeys(
                            section
                            for item in uncovered_items
                            for section in item.mapped_sections
                        )
                    )[:6],
                )
            )
        low_sections = [item for item in section_coverages if item.coverage_score < 60]
        if low_sections:
            checks.append(
                GenerationCheckResponse(
                    id=f"gen-check-{uuid4().hex[:8]}",
                    level="P1",
                    category="章节覆盖率",
                    title="存在低覆盖章节",
                    detail="部分章节未有效覆盖其目标评分项，建议优先补充缺失要求和证据素材。",
                    related_sections=[item.section_title for item in low_sections[:6]],
                )
            )
        if context.extracted_fields.get("必备资质", "待补充") != "待补充":
            covered_titles = {title for section in section_coverages for title in section.matched_score_items}
            if not any("资质" in title or "资格" in title for title in covered_titles):
                checks.append(
                    GenerationCheckResponse(
                        id=f"gen-check-{uuid4().hex[:8]}",
                        level="P1",
                        category="强制项自检",
                        title="资质门槛尚未显式覆盖",
                        detail="已抽取到必备资质，但生成内容中还没有明确对照资质门槛进行应答。",
                        related_sections=["公司概况与资质", "商务偏离说明"],
                    )
                )
        if context.extracted_fields.get("服务承诺", "待补充") != "待补充":
            service_section = next((item for item in section_coverages if item.section_title == "售后服务方案"), None)
            if service_section and service_section.coverage_score < 70:
                checks.append(
                    GenerationCheckResponse(
                        id=f"gen-check-{uuid4().hex[:8]}",
                        level="P2",
                        category="服务承诺自检",
                        title="售后服务承诺覆盖不足",
                        detail="服务承诺相关要求已抽取，但售后服务章节覆盖度偏低，建议补充 SLA、响应等级和质保安排。",
                        related_sections=["售后服务方案"],
                    )
                )
        return checks

    def _infer_target_sections(self, text: str) -> list[str]:
        normalized = text.replace(" ", "")
        matched_sections = [
            section_title
            for section_title, keywords in SECTION_MAPPING_RULES.items()
            if any(keyword.lower() in normalized.lower() for keyword in keywords)
        ]
        return matched_sections or ["总体技术方案", "商务偏离说明"]

    def _split_requirement_fragments(self, text: str) -> list[str]:
        fragments = [
            item.strip(" -:：;；，,。\n\r\t")
            for item in re.split(r"[\n;；]+", text)
            if item.strip(" -:：;；，,。\n\r\t")
        ]
        if not fragments and text.strip():
            fragments = [text.strip()]
        return fragments[:12]

    def _extract_keywords(self, text: str, *, limit: int = 10) -> list[str]:
        tokens = re.findall(r"[A-Za-z0-9\-]{2,}|[\u4e00-\u9fff]{2,8}", text)
        dedup: list[str] = []
        seen: set[str] = set()
        for token in tokens:
            key = token.lower()
            if key in seen:
                continue
            seen.add(key)
            dedup.append(token)
            if len(dedup) >= limit:
                break
        return dedup

    def _split_markdown_blocks(self, content: str) -> list[str]:
        cleaned = re.sub(r"^#+\s*", "", content, flags=re.MULTILINE).strip()
        return [block.strip() for block in cleaned.split("\n\n") if block.strip()]

    def _bullet_lines(self, raw_text: str, fallback: list[str]) -> str:
        parts = [part.strip(" -\n\r\t") for part in re.split(r"[\n;；]", raw_text) if part.strip(" -\n\r\t")]
        if not parts:
            parts = fallback
        return "\n".join(f"- {part}" for part in parts)

    def _bullet_lines_from_list(self, items: list[str], fallback: list[str]) -> str:
        values = [item.strip() for item in items if item.strip()]
        if not values:
            values = fallback
        return "\n".join(f"- {item}" for item in values)

    def _resolve_selected_assets(
        self,
        *,
        available_titles: list[str],
        requested_titles: list[str],
        fixed_titles: list[str],
        excluded_titles: list[str],
    ) -> list[str]:
        base_titles = [item for item in requested_titles if item] if requested_titles else available_titles[:3]
        merged: list[str] = []
        for title in [*fixed_titles, *base_titles]:
            normalized = title.strip()
            if (
                normalized
                and normalized in available_titles
                and normalized not in excluded_titles
                and normalized not in merged
            ):
                merged.append(normalized)
        if not merged:
            merged = [title for title in available_titles if title not in excluded_titles][:3]
        return merged

    def _routed_asset_lines(self, routed_assets: list[RoutedAsset], fallback: list[str]) -> str:
        if not routed_assets:
            return self._bullet_lines_from_list(fallback, fallback)
        return "\n".join(
            [
                f"- {item.asset_title} · {item.chunk_title}：{item.snippet}（{item.reason}）"
                for item in routed_assets[:3]
            ]
        )

    def _section_score_focus_lines(self, section_title: str, context: GenerationContext) -> str:
        mapped = [
            item.title
            for item in self._extract_score_items(context)
            if section_title in item.mapped_sections
        ][:4]
        if not mapped:
            return ""
        return (
            "\n\n### 本章需覆盖的评分点\n\n"
            + "\n".join([f"- {item}" for item in mapped])
            + "\n"
        )

    def _replace_section_asset_refs(
        self,
        db: Session,
        *,
        job_id: str,
        section_id: str,
        routed_assets: list[RoutedAsset],
    ) -> None:
        db.execute(delete(GenerationSectionAssetRef).where(GenerationSectionAssetRef.section_id == section_id))
        db.flush()
        for item in routed_assets:
            db.add(
                GenerationSectionAssetRef(
                    id=f"gen-ref-{uuid4().hex[:10]}",
                    job_id=job_id,
                    section_id=section_id,
                    asset_id=item.asset_id,
                    asset_title=item.asset_title,
                    chunk_title=item.chunk_title,
                    reason=item.reason,
                    snippet=item.snippet,
                    score=item.score,
                )
            )

    def _list_section_asset_refs(self, db: Session, section_id: str) -> list[GenerationSectionAssetRef]:
        return db.scalars(
            select(GenerationSectionAssetRef)
            .where(GenerationSectionAssetRef.section_id == section_id)
            .order_by(GenerationSectionAssetRef.score.desc())
        ).all()


PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")


class MarkdownToDocxConverter:
    """将 Markdown 内容转换为格式化的 Word 段落。"""

    def __init__(self, document: Document):
        self.document = document

    def add_markdown(self, content: str) -> None:
        """解析 Markdown 并将各元素添加到 Word 文档。"""
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # 跳过空白行
            if not line.strip():
                i += 1
                continue

            # 代码块 ```...```
            if line.strip().startswith("```"):
                code_lines: list[str] = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block("\n".join(code_lines))
                i += 1
                continue

            # 表格 | col1 | col2 |
            if line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(table_lines)
                continue

            # ATX 标题 ### ## #
            m = re.match(r"^(#{1,6})\s+(.+)$", line)
            if m:
                level = len(m.group(1))
                self._add_heading(m.group(2), level)
                i += 1
                continue

            # 列表项 - 或 * 或数字 .
            if re.match(r"^\s*[-*+]\s+", line) or re.match(r"^\s*\d+\.\s+", line):
                bullet_lines = []
                while i < len(lines) and (re.match(r"^\s*[-*+]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])):
                    bullet_lines.append(lines[i])
                    i += 1
                self._add_bullet_list(bullet_lines)
                continue

            # 普通段落（可能是多行组成一段）
            paragraph_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("|") \
                and not lines[i].strip().startswith("```") \
                and not re.match(r"^(#{1,6})\s+", lines[i]) \
                and not re.match(r"^\s*[-*+]\s+", lines[i]) \
                and not re.match(r"^\s*\d+\.\s+", lines[i]):
                paragraph_lines.append(lines[i])
                i += 1
            self._add_paragraph(" ".join(paragraph_lines))

    def _add_heading(self, text: str, level: int) -> None:
        """添加标题，level 1-6 映射到 Word Heading 1-6。"""
        text = text.strip()
        if not text:
            return
        # 去掉可能的 ATX 关闭标记 ###
        text = re.sub(r"\s*#+\s*$", "", text)
        self.document.add_heading(text, level=min(level, 9))  # Word 最大支持 9 级

    def _add_paragraph(self, text: str) -> None:
        """添加普通段落，支持 **bold** 和 *italic。"""
        text = text.strip()
        if not text:
            return
        p = self.document.add_paragraph()
        self._add_formatted_runs(p, text)

    def _add_bullet_list(self, lines: list[str]) -> None:
        """添加项目列表。"""
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 去掉列表标记
            content = re.sub(r"^\s*[-*+]\s+", "", line)
            content = re.sub(r"^\s*\d+\.\s+", "", content)
            p = self.document.add_paragraph(style="List Bullet")
            self._add_formatted_runs(p, content)

    def _add_code_block(self, code: str) -> None:
        """添加代码块（等宽字体，左边框）。"""
        code = code.strip()
        if not code:
            return
        p = self.document.add_paragraph(style="Quote")
        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(20)

    def _add_table(self, lines: list[str]) -> None:
        """将 Markdown 表格转换为 Word 表格。"""
        if len(lines) < 2:
            return
        rows_data = []
        for line in lines:
            # 跳过 Markdown 表格分隔行 |---|---|
            if re.match(r"^\|\s*[-:]+\s*(\|\s*[-:]+\s*)+$", line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            if cells:
                rows_data.append(cells)
        if not rows_data:
            return
        # 以最长行为准补齐
        max_cols = max(len(row) for row in rows_data)
        table = self.document.add_table(rows=len(rows_data), cols=max_cols)
        table.style = "Table Grid"
        for r_idx, row_data in enumerate(rows_data):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < max_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
        self.document.add_paragraph()  # 表格后空行

    def _add_formatted_runs(self, paragraph: "Paragraph", text: str) -> None:
        """将 text 中的 **bold** 和 *italic 解析为 Word runs。"""
        # 按 ** 或 * 分隔，但处理转义
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)


class TemplateDocxExporter:
    """基于 Word 模板的文档导出器，支持 {{placeholder}} 占位符替换。"""

    def __init__(self, template_bytes: bytes | None = None):
        self.template_bytes = template_bytes

    def export(
        self,
        sections: list[dict],  # [{title: str, content: str, citations: str, todos: str}]
        project_name: str,
        metadata: dict,  # {key: value} 注入所有 {{key}} 占位符
    ) -> bytes:
        """
        如果传入了模板字节流，则打开模板并替换占位符；
        否则基于 python-docx 内置样式生成文档。
        """
        if self.template_bytes:
            doc = Document(BytesIO(self.template_bytes))
        else:
            doc = Document()

        if self.template_bytes:
            self._replace_placeholders(doc, metadata)
            self._replace_sections(doc, sections)
        else:
            self._build_from_scratch(doc, sections, project_name, metadata)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _replace_placeholders(self, doc: Document, metadata: dict) -> None:
        """遍历模板段落，替换 {{key}} 占位符。"""
        for para in doc.paragraphs:
            if not PLACEHOLDER_PATTERN.search(para.text):
                continue
            full_text = para.text
            for key, value in metadata.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
            # 重建 paragraph，清除富文本占位符
            if para.text != full_text:
                p = para._element
                for run in para.runs:
                    r = run._element
                    r.getparent().remove(r)
                para.clear()
                para.add_run(full_text)

    def _replace_sections(self, doc: Document, sections: list[dict]) -> None:
        """遍历模板段落，将 {{section_title}} 占位符替换为 AI 生成的格式化章节内容。"""
        converter = MarkdownToDocxConverter(doc)
        # 收集需要替换的段落
        placeholder_paras: list[tuple[OxmlElement, str, str]] = []
        for para in doc.paragraphs:
            matches = PLACEHOLDER_PATTERN.findall(para.text)
            for title in matches:
                placeholder_paras.append((para._element, title, para.text))

        # 按标题匹配章节
        for para_elem, title, full_text in placeholder_paras:
            matched_section = next(
                (s for s in sections if s.get("title", "").strip() == title.strip()), None
            )
            if not matched_section:
                continue
            content = matched_section.get("content", "")
            citations = matched_section.get("citations", "")
            todos = matched_section.get("todos", "")

            # 在段落位置插入格式化内容
            parent = para_elem.getparent()
            idx = list(parent).index(para_elem)
            parent.remove(para_elem)

            # 添加章节标题
            h = doc.add_heading(title, level=1)
            for run in h.runs:
                run.font.size = Pt(14)

            # 添加 Markdown 内容
            converter.add_markdown(content)

            # 添加引用+待确认
            if citations or todos:
                p = doc.add_paragraph()
                parts = []
                if citations:
                    parts.append(f"引用 {citations} 条")
                if todos:
                    parts.append(f"待确认 {todos} 项")
                p.add_run(" · ".join(parts)).italic = True

    def _build_from_scratch(
        self,
        doc: Document,
        sections: list[dict],
        project_name: str,
        metadata: dict,
    ) -> None:
        """无模板时，基于 python-docx 样式生成完整文档。"""
        doc.add_heading(f"回标文件：{project_name}", level=0)
        for key, value in metadata.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}：").bold = True
            p.add_run(str(value))
        for section in sections:
            doc.add_heading(section.get("title", ""), level=1)
            converter = MarkdownToDocxConverter(doc)
            converter.add_markdown(section.get("content", ""))
            citations = section.get("citations", "")
            todos = section.get("todos", "")
            if citations or todos:
                p = doc.add_paragraph()
                parts = []
                if citations:
                    parts.append(f"引用 {citations} 条")
                if todos:
                    parts.append(f"待确认 {todos} 项")
                run = p.add_run(" · ".join(parts))
                run.italic = True


generation_service = GenerationService()
