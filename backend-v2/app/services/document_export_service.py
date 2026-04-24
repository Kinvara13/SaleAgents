"""Document export service: convert editable_content to Word/Excel/PDF files."""

import os
import re
from datetime import datetime
from pathlib import Path

from fastapi import HTTPException, status
from sqlalchemy.orm import Session

from app.models.business_document import BusinessDocument
from app.models.technical_document import TechnicalDocument

# Export base directory (relative to project root)
EXPORTS_DIR = Path(__file__).resolve().parent.parent.parent / "exports"


def _ensure_exports_dir(project_id: str) -> Path:
    d = EXPORTS_DIR / project_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _is_table_like(doc_type: str, doc_name: str, content: str) -> bool:
    """Heuristic to decide whether a document is table-like and should default to Excel."""
    table_keywords = ["偏离表", "申报表", "确认单", "清单", "汇总", "业绩", "著作权", "自查"]
    name_or_type = f"{doc_name} {doc_type}".lower()
    if any(k in name_or_type for k in table_keywords):
        return True
    # If content has markdown tables
    if re.search(r"\|.*\|.*\|", content):
        return True
    return False


def _extract_markdown_tables(content: str) -> list[list[list[str]]]:
    """Extract all markdown tables from content. Returns list of tables (list of rows)."""
    tables = []
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "|" in line:
            # Potential table start
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            # Filter out separator lines like |---|---|
            rows = []
            for tl in table_lines:
                if re.match(r"^\|?[\-\s|:]+\|", tl):
                    continue
                cells = [c.strip() for c in tl.split("|")]
                # Remove empty first/last cells caused by leading/trailing |
                if cells and cells[0] == "":
                    cells = cells[1:]
                if cells and cells[-1] == "":
                    cells = cells[:-1]
                if cells:
                    rows.append(cells)
            if rows:
                tables.append(rows)
            continue
        i += 1
    return tables


def _content_to_docx(content: str, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(10.5)

    lines = content.splitlines()
    in_table = False
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_buffer
        if not table_buffer:
            in_table = False
            return
        # Determine max columns
        max_cols = max(len(r) for r in table_buffer) if table_buffer else 0
        table = doc.add_table(rows=len(table_buffer), cols=max_cols)
        table.style = "Table Grid"
        for ri, row_cells in enumerate(table_buffer):
            row = table.rows[ri]
            for ci, val in enumerate(row_cells):
                if ci < max_cols:
                    row.cells[ci].text = val
        doc.add_paragraph()
        table_buffer = []
        in_table = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_table:
                flush_table()
            continue

        # Header
        if stripped.startswith("# "):
            if in_table:
                flush_table()
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue
        if stripped.startswith("## "):
            if in_table:
                flush_table()
            p = doc.add_heading(stripped[3:], level=2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue
        if stripped.startswith("### "):
            if in_table:
                flush_table()
            p = doc.add_heading(stripped[4:], level=3)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        # Table row
        if "|" in stripped:
            cells = [c.strip() for c in stripped.split("|")]
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]
            # Skip separator lines
            if all(re.match(r"^[\-\s:]+$", c) for c in cells if c):
                continue
            if cells:
                in_table = True
                table_buffer.append(cells)
            continue
        else:
            if in_table:
                flush_table()

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(text, style="List Number")
            continue

        # Normal paragraph
        doc.add_paragraph(stripped)

    if in_table:
        flush_table()

    doc.save(str(output_path))


def _content_to_xlsx(content: str, output_path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    wb = Workbook()
    ws = wb.active
    if ws is None:
        ws = wb.create_sheet("Sheet1")

    tables = _extract_markdown_tables(content)
    if tables:
        # Use the largest table as main sheet data, prepend non-table text as header rows
        main_table = max(tables, key=lambda t: len(t) * len(t[0]) if t else 0)
    else:
        main_table = []

    # Write non-table content as header comments above the table
    header_lines = []
    lines = content.splitlines()
    for line in lines:
        stripped = line.strip()
        if "|" in stripped and re.search(r"\|.*\|", stripped):
            break
        if stripped:
            header_lines.append(stripped)

    row_idx = 1
    for hl in header_lines:
        cell = ws.cell(row=row_idx, column=1, value=hl)
        cell.font = Font(name="宋体", size=11, bold=True)
        row_idx += 1

    if main_table:
        row_idx += 1  # blank row
        header_fill = PatternFill(start_color="B4C7DC", end_color="B4C7DC", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        for ri, row_cells in enumerate(main_table):
            for ci, val in enumerate(row_cells):
                cell = ws.cell(row=row_idx, column=ci + 1, value=val)
                cell.border = thin_border
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.font = Font(name="宋体", size=10)
                if ri == 0:
                    cell.font = Font(name="宋体", size=10, bold=True)
                    cell.fill = header_fill
            row_idx += 1

    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[col_letter].width = adjusted_width

    wb.save(str(output_path))


def export_document(
    db: Session,
    project_id: str,
    doc_id: str,
    doc_kind: str,  # "business" or "technical"
    fmt: str | None = None,
) -> dict:
    """Export a document to a file and return download URL info."""
    if doc_kind == "business":
        doc = db.query(BusinessDocument).filter(
            BusinessDocument.id == doc_id,
            BusinessDocument.project_id == project_id,
        ).first()
    else:
        doc = db.query(TechnicalDocument).filter(
            TechnicalDocument.id == doc_id,
            TechnicalDocument.project_id == project_id,
        ).first()

    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="文档不存在")

    content = doc.editable_content or doc.original_content or ""
    if not content.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="文档内容为空，无法导出")

    # Determine format
    if fmt not in ("docx", "xlsx"):
        if _is_table_like(doc.doc_type, doc.doc_name, content):
            fmt = "xlsx"
        else:
            fmt = "docx"

    out_dir = _ensure_exports_dir(project_id)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_doc_type = re.sub(r"[^\w\-]", "_", doc.doc_type)
    filename = f"{safe_doc_type}_{ts}.{fmt}"
    output_path = out_dir / filename

    if fmt == "docx":
        _content_to_docx(content, output_path)
    else:
        _content_to_xlsx(content, output_path)

    # Return a relative URL that can be served via static files
    download_url = f"/exports/{project_id}/{filename}"
    return {
        "download_url": download_url,
        "filename": filename,
        "format": fmt,
        "file_path": str(output_path),
    }
