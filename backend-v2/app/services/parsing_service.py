import logging
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from uuid import uuid4

from sqlalchemy.orm import Session

from app.models.parsing_section import ParsingSection
from app.models.project import Project
from app.schemas.parsing import ParsingSectionSummary, ParsingSectionDetail, ParsingSectionUpdateRequest
from app.schemas.workspace import ParseSection
from app.services.llm_parsing_client import llm_parsing_client
from app.services.rule_based_extractor import rule_extractor

logger = logging.getLogger(__name__)

BUSINESS_SECTIONS = [
    "商务偏离表",
    "应答承诺函",
    "授权委托书",
    "营业执照",
    "资格审查资料",
    "应答保证金",
    "封面",
]

TECH_SECTIONS = [
    "技术条款偏离表",
    "CMMI证书",
    "计算机软件著作权证书",
    "项目案例",
    "自查确认单",
    "服务承诺书",
    "封面",
]

# Chunking constants
MAX_CHUNK_CHARS = 12000
SUMMARIZE_THRESHOLD = 15000
MIN_CHUNK_CHARS = 200
MAX_STAR_ITEMS = 20

FIELD_LABELS = [
    "项目名称", "招标编号", "标书类型", "投标截止时间", "预算金额",
    "标书起始时间", "标书结束时间", "是否有保证金", "保证金金额", "保证金形式",
    "必备资质", "付款条款", "交付周期", "评分重点", "技术要求", "服务承诺",
    "是否需要签字盖章", "是否有项目澄清会", "项目澄清会时间", "项目澄清会链接",
    "废标条款", "合同条款", "商务条款", "投标人须知",
]

# Heading patterns for Chinese tender documents
HEADING_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百\d]+章[、\.．\s]+"),
    re.compile(r"^第[一二三四五六七八九十\d]+节[、\.．\s]+"),
    re.compile(r"^[一二三四五六七八九十]+[、\.．]\s*"),
    re.compile(r"^\([一二三四五六七八九十\d]+\)"),
    re.compile(r"^\(\d+\)"),
    re.compile(r"^\d+[\.．]\d+\s+"),
    re.compile(r"^\d+[、\.．]\s+\S"),
    re.compile(r"^(附件|附录)\d*\s*"),
]

KEYWORD_SECTION_HEADINGS = {
    "招标公告", "投标人须知", "评标办法", "合同条款", "技术要求",
    "商务要求", "资格审查", "投标文件格式", "报价要求", "服务承诺",
    "质保要求", "交付要求", "项目背景", "项目概况", "建设目标",
    "总体要求", "功能需求", "非功能需求", "安全要求", "附件",
}


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    for pattern in HEADING_PATTERNS:
        if pattern.match(stripped):
            return True
    # Heuristic: short line ending with punctuation and containing section keywords
    if len(stripped) < 40:
        clean = stripped.rstrip("：:. \n")
        if clean in KEYWORD_SECTION_HEADINGS:
            return True
        if any(kw in clean for kw in ["说明", "要求", "条款", "标准", "规范", "评审", "评分", "合同", "资质", "附录"]):
            if stripped.endswith(("：", ":", "。")):
                return True
    return False


def _normalize_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text or "")
    text = re.sub(r"[ \t\u3000]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_docx_text(file_path: Path) -> str:
    from docx import Document

    doc = Document(str(file_path))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            lines.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_office_text_with_textutil(file_path: Path, ext: str) -> str:
    if ext not in {".doc", ".rtf"}:
        return ""
    try:
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / f"{file_path.stem}.txt"
            subprocess.run(
                ["textutil", "-convert", "txt", "-output", str(out_path), str(file_path)],
                check=True,
                capture_output=True,
                timeout=60,
            )
            return out_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.warning("textutil extraction failed for %s: %s", file_path, exc)
        return ""


def _extract_excel_text(file_path: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(str(file_path), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        lines.append(f"【工作表：{sheet}】")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_pdf_pages(file_path: Path) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []

    try:
        import fitz

        with fitz.open(str(file_path)) as doc:
            for i, page in enumerate(doc, start=1):
                text = page.get_text("text") or ""
                if text.strip():
                    pages.append((i, text))
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed for %s: %s", file_path, exc)

    if pages and sum(len(text.strip()) for _, text in pages) >= 80:
        return pages

    try:
        import pypdf

        reader = pypdf.PdfReader(str(file_path))
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append((i, text))
    except Exception as exc:
        logger.warning("pypdf extraction failed for %s: %s", file_path, exc)

    if pages and sum(len(text.strip()) for _, text in pages) >= 80:
        return pages

    return _ocr_pdf_pages(file_path)


def _ocr_pdf_pages(file_path: Path) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []
    try:
        import fitz

        try:
            from ocrmac import ocrmac

            with fitz.open(str(file_path)) as doc:
                for i, page in enumerate(doc, start=1):
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    with tempfile.NamedTemporaryFile(suffix=".png") as image_file:
                        image_file.write(pix.tobytes("png"))
                        image_file.flush()
                        result = ocrmac.OCR(str(image_file.name), language_preference=["zh-Hans", "en"]).recognize()
                    text = "\n".join(item[0] for item in result if item and item[0]).strip()
                    if text:
                        pages.append((i, text))
            if pages:
                logger.info("OCR extracted %d pages from %s using ocrmac", len(pages), file_path)
                return pages
        except Exception as exc:
            logger.warning("ocrmac extraction failed for %s: %s", file_path, exc)

        try:
            from rapidocr_onnxruntime import RapidOCR

            engine = RapidOCR()
            with fitz.open(str(file_path)) as doc:
                for i, page in enumerate(doc, start=1):
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    with tempfile.NamedTemporaryFile(suffix=".png") as image_file:
                        image_file.write(pix.tobytes("png"))
                        image_file.flush()
                        result, _ = engine(str(image_file.name))
                    text = "\n".join(row[1] for row in (result or []) if len(row) > 1 and row[1]).strip()
                    if text:
                        pages.append((i, text))
            if pages:
                logger.info("OCR extracted %d pages from %s using rapidocr", len(pages), file_path)
        except Exception as exc:
            logger.warning("RapidOCR extraction failed for %s: %s", file_path, exc)
    except Exception as exc:
        logger.warning("PDF OCR setup failed for %s: %s", file_path, exc)

    return pages


def _extract_text_from_file(file_path: Path, ext: str) -> tuple[str, list[tuple[int, str]]]:
    """Extract text from a file. Returns (full_text, list_of(page_num, page_text))."""
    pages: list[tuple[int, str]] = []
    full_text = ""

    try:
        if ext == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="replace")
            pages.append((1, _normalize_text(text)))
            full_text = pages[0][1]
        elif ext == ".docx":
            try:
                text = _normalize_text(_extract_docx_text(file_path))
                pages.append((1, text))
                full_text = text
            except Exception as exc:
                logger.warning("DOCX extraction failed for %s: %s", file_path, exc)
        elif ext == ".doc":
            text = _normalize_text(_extract_office_text_with_textutil(file_path, ext))
            if text:
                pages.append((1, text))
                full_text = text
        elif ext == ".xlsx":
            try:
                text = _normalize_text(_extract_excel_text(file_path))
                pages.append((1, text))
                full_text = text
            except Exception as exc:
                logger.warning("XLSX extraction failed for %s: %s", file_path, exc)
        elif ext == ".xls":
            logger.warning("Legacy XLS extraction is unavailable without xlrd/libreoffice: %s", file_path)
        elif ext == ".pdf":
            pages = [(p, _normalize_text(text)) for p, text in _extract_pdf_pages(file_path) if text.strip()]
            full_text = "\n\n".join(f"--- Page {p[0]} ---\n{p[1]}" for p in pages)
        else:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            pages.append((1, _normalize_text(text)))
            full_text = pages[0][1]
    except Exception as e:
        logger.error("Text extraction failed for %s: %s", file_path, e)
        full_text = f"[文件读取失败: {e}]"

    return full_text, pages


def _chunk_by_headings(pages: list[tuple[int, str]]) -> list[dict[str, Any]]:
    """Group page text into logical chunks by heading detection."""
    chunks: list[dict[str, Any]] = []
    current_title = "文档开头"
    current_lines: list[str] = []
    current_start_page = 1

    for page_num, page_text in pages:
        lines = page_text.splitlines()
        for line in lines:
            if _is_heading(line):
                # Save previous chunk
                if current_lines:
                    content = "\n".join(current_lines).strip()
                    if len(content) >= MIN_CHUNK_CHARS:
                        chunks.append({
                            "title": current_title,
                            "start_page": current_start_page,
                            "end_page": page_num,
                            "content": content,
                        })
                current_title = line.strip().rstrip("：:. \n")
                current_lines = []
                current_start_page = page_num
            else:
                current_lines.append(line)

    # Save last chunk
    if current_lines:
        content = "\n".join(current_lines).strip()
        if len(content) >= MIN_CHUNK_CHARS:
            chunks.append({
                "title": current_title,
                "start_page": current_start_page,
                "end_page": pages[-1][0] if pages else current_start_page,
                "content": content,
            })

    return chunks


def _split_oversized_chunks(chunks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Split chunks that exceed MAX_CHUNK_CHARS into smaller pieces."""
    result: list[dict[str, Any]] = []
    for chunk in chunks:
        content = chunk["content"]
        if len(content) <= MAX_CHUNK_CHARS:
            result.append(chunk)
            continue
        # Try to split by paragraphs
        paragraphs = content.split("\n\n")
        parts: list[str] = []
        current_part = ""
        for para in paragraphs:
            if len(current_part) + len(para) + 2 > MAX_CHUNK_CHARS:
                if current_part:
                    parts.append(current_part.strip())
                current_part = para
            else:
                current_part = current_part + "\n\n" + para if current_part else para
        if current_part:
            parts.append(current_part.strip())

        for idx, part in enumerate(parts, start=1):
            if len(part) >= MIN_CHUNK_CHARS:
                result.append({
                    "title": f"{chunk['title']} (Part {idx})" if len(parts) > 1 else chunk["title"],
                    "start_page": chunk["start_page"],
                    "end_page": chunk["end_page"],
                    "content": part,
                })
    return result


def _summarize_chunk(content: str, section_name: str) -> str:
    """Use LLM to summarize a long chunk. Falls back to tail-aware truncation."""
    try:
        summary = llm_parsing_client.summarize_text(content, title=section_name, max_words=2000)
        if summary:
            return summary
    except Exception as e:
        logger.warning("LLM summarization failed for chunk '%s': %s", section_name, e)

    # Fallback: smart truncation preserving head and tail
    head_len = 6000
    tail_len = 4000
    if len(content) > head_len + tail_len + 200:
        return (
            content[:head_len]
            + "\n\n...[中间内容省略]...\n\n"
            + content[-tail_len:]
        )
    return content[:MAX_CHUNK_CHARS]


@dataclass
class ParsingContext:
    parse_sections: list = field(default_factory=list)
    source_excerpt: str = ""


class ParsingService:
    """Real PDF/document parsing with chunking and summarization."""

    # --- Backward-compatible module-level API wrappers ---

    def list_sections(self, db: Session, project_id: str) -> list[ParsingSectionSummary]:
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            from fastapi import HTTPException, status
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="项目不存在")
        sections = (
            db.query(ParsingSection)
            .filter(ParsingSection.project_id == project_id)
            .order_by(ParsingSection.section_type, ParsingSection.id)
            .all()
        )
        return [ParsingSectionSummary.model_validate(s) for s in sections]

    def get_section_detail(self, db: Session, project_id: str, section_id: str) -> ParsingSectionDetail:
        section = (
            db.query(ParsingSection)
            .filter(ParsingSection.id == section_id, ParsingSection.project_id == project_id)
            .first()
        )
        if not section:
            from fastapi import HTTPException, status
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="章节不存在")
        return ParsingSectionDetail.model_validate(section)

    def update_section(
        self, db: Session, project_id: str, section_id: str, payload: ParsingSectionUpdateRequest
    ) -> ParsingSectionDetail:
        section = (
            db.query(ParsingSection)
            .filter(ParsingSection.id == section_id, ParsingSection.project_id == project_id)
            .first()
        )
        if not section:
            from fastapi import HTTPException, status
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="章节不存在")
        if payload.content is not None:
            section.content = payload.content
        db.commit()
        db.refresh(section)
        return ParsingSectionDetail.model_validate(section)

    # --- Core parsing logic ---

    def parse_document(
        self,
        db: Session,
        project_id: str,
        file_path: Path,
        filename: str,
        clear_existing: bool = True,
    ) -> list[ParsingSectionSummary]:
        """Parse a single document (PDF/DOCX/TXT/etc.) into ParsingSections."""
        ext = Path(filename).suffix.lower()
        full_text, pages = _extract_text_from_file(file_path, ext)

        if not full_text or len(full_text.strip()) < 50:
            logger.warning("No text extracted from %s", filename)
            return self._create_placeholder_sections(db, project_id, filename, clear_existing=clear_existing)

        # Rule-based pre-extraction (fast, deterministic)
        rule_results = rule_extractor.extract_fields(full_text)
        logger.info(f"[RuleExtractor] Extracted fields: {list(rule_results.keys())}, star_items: {len(rule_results.get('star_items', []))}")

        # 1. Chunk by headings
        chunks = _chunk_by_headings(pages)

        # 2. Split oversized chunks
        chunks = _split_oversized_chunks(chunks)

        # 4. Merge tiny chunks with neighbors
        chunks = self._merge_tiny_chunks(chunks)

        # 5. Extract tender fields via LLM using curated local text only.
        curated_text = self._curate_text_for_llm(full_text)
        tender_info = llm_parsing_client.extract_tender_fields(curated_text)
        if not tender_info:
            logger.warning("No LLM tender fields extracted for %s; using local rules only", filename)

        # Merge: rule-based takes priority for structured fields, LLM for complex fields
        merged_info = self._merge_extraction_results(rule_results, tender_info or {}, full_text)
        self._fill_from_filename(merged_info, filename)

        # 6. Persist
        if clear_existing:
            db.query(ParsingSection).filter(ParsingSection.project_id == project_id).delete()

        created: list[ParsingSection] = []

        # Document structure chunks
        for chunk in chunks:
            section = ParsingSection(
                id=f"sec_{uuid4().hex[:12]}",
                project_id=project_id,
                section_name=chunk["title"],
                section_type="评审" if any(k in chunk["title"] for k in ["评分", "评审", "招标"]) else "内容",
                content=chunk["content"],
                is_star_item=any(k in chunk["title"] for k in ["★", "星标", "废标", "否决", "红线", "关键条款", "实质性要求"]),
                source_file=filename,
            )
            db.add(section)
            created.append(section)

        # Key-field sections from merged extraction
        if merged_info:
            scoring_text = self._extract_field_value(merged_info, "评分重点")
            if scoring_text and len(scoring_text) > 20:
                created.append(self._add_section(
                    db, project_id, "评分规则解析", "评审", scoring_text, True, filename
                ))

            tech_req = self._extract_field_value(merged_info, "技术要求")
            if tech_req and len(tech_req) > 20:
                created.append(self._add_section(
                    db, project_id, "技术要求", "评审", tech_req, True, filename
                ))

            qual_req = self._extract_field_value(merged_info, "必备资质")
            if qual_req and len(qual_req) > 20:
                created.append(self._add_section(
                    db, project_id, "资质要求", "评审", qual_req, True, filename
                ))

            # New fields from rule-based extraction
            for field_name in ["废标条款", "合同条款", "商务条款", "投标人须知"]:
                field_val = self._extract_field_value(merged_info, field_name)
                if field_val and len(field_val) > 20:
                    created.append(self._add_section(
                        db, project_id, field_name, "评审", field_val, True, filename
                    ))

            # Extract star items from merged response
            star_items_list = merged_info.get("星标项列表", [])
            if isinstance(star_items_list, list):
                for star_item in star_items_list[:MAX_STAR_ITEMS]:
                    if isinstance(star_item, dict):
                        star_name = star_item.get("name", "")
                        star_content = star_item.get("content", "")
                        if star_name and star_content:
                            created.append(self._add_section(
                                db, project_id, star_name, "评审", star_content, True, filename
                            ))

            # Update project extracted_fields
            project = db.query(Project).filter(Project.id == project_id).first()
            if project:
                extracted = self._merge_project_extracted_fields(project.extracted_fields or [], merged_info)
                project.extracted_fields = extracted

        # Placeholder business/tech sections expected by the frontend
        for name in BUSINESS_SECTIONS:
            created.append(self._add_section(
                db, project_id, name, "商务",
                f"【{name}】\n\n解析自招标文件 {filename}，请根据招标要求及公司实际情况填写。",
                False, filename
            ))
        for name in TECH_SECTIONS:
            created.append(self._add_section(
                db, project_id, name, "技术",
                f"【{name}】\n\n解析自招标文件 {filename}，请根据技术规范书要求填写。",
                False, filename
            ))

        db.commit()
        return [ParsingSectionSummary.model_validate(s) for s in created]

    def _merge_project_extracted_fields(self, existing_fields: list | dict, merged_info: dict) -> list[dict[str, str]]:
        existing: dict[str, dict[str, str]] = {}
        if isinstance(existing_fields, list):
            for item in existing_fields:
                if isinstance(item, dict) and item.get("label"):
                    existing[str(item["label"])] = {
                        "label": str(item["label"]),
                        "value": str(item.get("value", "")),
                        "confidence": str(item.get("confidence", "70%")),
                    }

        for key in FIELD_LABELS:
            val = self._extract_field_value(merged_info, key)
            if self._is_useful_value(val):
                confidence = self._extract_confidence(merged_info, key) or "85%"
                current = existing.get(key, {})
                if not self._is_useful_value(current.get("value")) or self._confidence_score(confidence) >= self._confidence_score(current.get("confidence", "")):
                    existing[key] = {"label": key, "value": val, "confidence": confidence}

        star_items_list = merged_info.get("星标项列表", [])
        if star_items_list and isinstance(star_items_list, list):
            star_summary = "; ".join([
                f"{item.get('name', '星标项')}: {item.get('content', '')[:100]}"
                for item in star_items_list[:MAX_STAR_ITEMS] if isinstance(item, dict)
            ])
            if star_summary:
                current = existing.get("星标项", {})
                merged_star = star_summary
                if self._is_useful_value(current.get("value")) and star_summary not in current.get("value", ""):
                    merged_star = f"{current['value']}; {star_summary}"
                existing["星标项"] = {"label": "星标项", "value": merged_star[:3000], "confidence": "90%"}

        return [existing[key] for key in [*FIELD_LABELS, "星标项"] if key in existing]

    @staticmethod
    def _confidence_score(confidence: str) -> int:
        match = re.search(r"\d+", str(confidence or ""))
        return int(match.group(0)) if match else 0

    def _fill_from_filename(self, merged_info: dict, filename: str) -> None:
        if self._is_useful_value(self._extract_field_value(merged_info, "项目名称")):
            return
        stem = Path(filename).stem
        stem = re.sub(r"^\d+[-_]", "", stem)
        stem = re.sub(r"[-_ ]?\d{8,14}$", "", stem)
        stem = re.sub(r"(采购合同|招标文件|采购文件|应答文件|投标文件)$", "", stem).strip("-_ ")
        if len(stem) >= 8 and re.search(r"[\u4e00-\u9fa5]", stem):
            merged_info["项目名称"] = {"value": stem, "confidence": "60%"}

    def _deprecated_noop(self) -> None:
        return None

    def _add_section(
        self, db: Session, project_id: str, name: str, sec_type: str,
        content: str, is_star: bool, filename: str
    ) -> ParsingSection:
        section = ParsingSection(
            id=f"sec_{uuid4().hex[:12]}",
            project_id=project_id,
            section_name=name,
            section_type=sec_type,
            content=content,
            is_star_item=is_star,
            source_file=filename,
        )
        db.add(section)
        return section

    def _extract_field_value(self, tender_info: dict, key: str) -> str:
        item = tender_info.get(key)
        if isinstance(item, dict):
            return str(item.get("value", ""))
        return str(item) if item else ""

    def _extract_confidence(self, tender_info: dict, key: str) -> str:
        item = tender_info.get(key)
        if isinstance(item, dict):
            return str(item.get("confidence", ""))
        return ""

    def _merge_extraction_results(self, rule_results: dict, llm_results: dict, full_text: str = "") -> dict:
        """Merge rule-based and LLM extraction results. Rule-based takes priority for structured fields."""
        merged = self._normalize_llm_results(llm_results or {})

        confidence = rule_results.get("confidence", {}) if isinstance(rule_results, dict) else {}

        structured_fields = {
            "project_name": "项目名称",
            "bid_number": "招标编号",
            "deadline": "投标截止时间",
        }
        for rule_key, label in structured_fields.items():
            rule_val = rule_results.get(rule_key)
            if self._is_useful_value(rule_val) and float(confidence.get(rule_key, 0) or 0) >= 0.4:
                merged[label] = {
                    "value": str(rule_val).strip(),
                    "confidence": f"{int(float(confidence.get(rule_key, 0.8)) * 100)}%",
                }

        # Complex fields: use rule_results section text if available
        complex_fields = {
            "scoring_criteria": "评分重点",
            "technical_requirements": "技术要求",
            "qualification_requirements": "必备资质",
            "disqualification_clauses": "废标条款",
            "contract_terms": "合同条款",
            "commercial_terms": "商务条款",
            "bidder_instructions": "投标人须知",
        }
        for key, display_name in complex_fields.items():
            rule_val = rule_results.get(key)
            if rule_val and isinstance(rule_val, str) and len(rule_val.strip()) > 8:
                merged[display_name] = {"value": rule_val, "confidence": f"{int(float(confidence.get(key, 0.8)) * 100)}%"}
            elif rule_val and isinstance(rule_val, dict) and rule_val.get("content"):
                merged[display_name] = {"value": rule_val["content"], "confidence": "80%"}

        for key, value in self._extract_local_basic_fields(full_text).items():
            if self._is_useful_value(value) and not self._is_useful_value(self._extract_field_value(merged, key)):
                merged[key] = {"value": value, "confidence": "70%"}

        # Star items: combine both sources, deduplicate by content similarity
        rule_star_items = rule_results.get("star_items", [])
        llm_star_items = llm_results.get("星标项列表", []) if llm_results else []

        combined_stars = []
        seen_contents = set()

        # Add rule-based star items first (higher priority)
        for item in rule_star_items:
            if isinstance(item, dict):
                content = item.get("content", "")
                name = item.get("name", "")
            elif isinstance(item, str):
                content = item
                name = "星标项"
            else:
                continue
            # Simple dedup by first 50 chars
            content_key = content[:50].strip()
            if content_key and content_key not in seen_contents:
                seen_contents.add(content_key)
                combined_stars.append({"name": name, "content": content})

        # Add LLM star items as supplement
        for item in llm_star_items:
            if isinstance(item, dict):
                content = item.get("content", "")
                name = item.get("name", "")
            else:
                continue
            content_key = content[:50].strip()
            if content_key and content_key not in seen_contents:
                seen_contents.add(content_key)
                combined_stars.append({"name": name, "content": content})

        if combined_stars:
            merged["星标项列表"] = combined_stars[:MAX_STAR_ITEMS]

        return merged

    def _normalize_llm_results(self, llm_results: dict) -> dict:
        merged: dict[str, Any] = {}
        aliases = {
            "评分规则": "评分重点",
            "项目预算": "预算金额",
            "截止时间": "投标截止时间",
            "资质要求": "必备资质",
        }
        for key, value in llm_results.items():
            label = aliases.get(key, key)
            if label == "星标项列表":
                merged[label] = value
                continue
            if label in FIELD_LABELS:
                if isinstance(value, dict):
                    raw = str(value.get("value", "")).strip()
                    if self._is_useful_value(raw):
                        merged[label] = {
                            "value": raw,
                            "confidence": str(value.get("confidence", "80%")),
                        }
                elif self._is_useful_value(value):
                    merged[label] = {"value": str(value).strip(), "confidence": "75%"}
        return merged

    def _extract_local_basic_fields(self, text: str) -> dict[str, str]:
        result: dict[str, str] = {}

        bid_type_patterns = [
            (r"(公开招标|邀请招标|竞争性谈判|竞争性磋商|询价采购|单一来源)", 1),
            (r"(?:采购方式|招标方式|标书类型)\s*[：:]\s*([^\n；;，, ]{2,20})", 1),
        ]
        for pattern, group in bid_type_patterns:
            m = re.search(pattern, text)
            if m:
                result["标书类型"] = m.group(group).strip()
                break

        budget_patterns = [
            r"(?:预算金额|采购预算|项目预算|最高限价|控制价)\s*[：:]?\s*((?:人民币|￥|¥)?\s*[\d,，.]+(?:\s*[万亿]?元)?)",
            r"((?:人民币|￥|¥)?\s*[\d,，.]+(?:\s*[万亿]?元)?)\s*(?:预算|最高限价|控制价)",
        ]
        for pattern in budget_patterns:
            m = re.search(pattern, text)
            if m:
                result["预算金额"] = re.sub(r"\s+", "", m.group(1))
                break

        sale_range = self._extract_date_range(
            text,
            ["招标文件获取时间", "获取招标文件", "发售时间", "购买标书时间", "报名时间"],
        )
        if sale_range:
            result["标书起始时间"], result["标书结束时间"] = sale_range

        guarantee_context = self._find_context(text, ["投标保证金", "保证金"])
        if guarantee_context:
            result["是否有保证金"] = "否" if re.search(r"不收取|无需|不要求|免收", guarantee_context) else "是"
            guarantee_idx = guarantee_context.find("保证金")
            amount_scope = guarantee_context[max(0, guarantee_idx):] if guarantee_idx >= 0 else guarantee_context
            amount_match = re.search(r"((?:人民币|￥|¥)?\s*[\d,，.]+(?:\s*[万亿]?元))", amount_scope)
            if amount_match:
                result["保证金金额"] = re.sub(r"\s+", "", amount_match.group(1))
            form_match = re.search(r"(银行转账|电汇|保函|银行保函|现金|支票|汇票|担保)", guarantee_context)
            if form_match:
                result["保证金形式"] = form_match.group(1)
        elif "保证金" not in text[:20000]:
            result["是否有保证金"] = "否"

        if re.search(r"(签字|签章|盖章|公章|电子章)", text[:50000]):
            result["是否需要签字盖章"] = "是"

        clarification = self._find_context(text, ["澄清会", "答疑会", "投标预备会"])
        if clarification:
            result["是否有项目澄清会"] = "否" if re.search(r"不召开|无|不组织", clarification) else "是"
            date = self._extract_first_datetime(clarification)
            if date:
                result["项目澄清会时间"] = date
            link = re.search(r"(https?://\S+|腾讯会议[：:]?\S+|会议号[：:]?\s*\S+)", clarification)
            if link:
                result["项目澄清会链接"] = link.group(1).strip()
        else:
            result["是否有项目澄清会"] = "否"

        delivery = self._find_context(text, ["服务期", "工期", "交付期", "交付周期", "合同履行期限"])
        if delivery:
            m = re.search(r"((?:合同签订|项目启动|中标通知书发出|采购人通知)?[^\n。；;]{0,20}(?:\d+\s*(?:个)?(?:日历日|工作日|天|个月|年)|至\s*\d{4}年\d{1,2}月\d{1,2}日)[^\n。；;]{0,30})", delivery)
            if m:
                result["交付周期"] = m.group(1).strip()

        payment = self._find_context(text, ["付款", "支付", "结算"])
        if payment and len(payment) > 20:
            result["付款条款"] = payment[:800]

        service = self._find_context(text, ["售后服务", "服务承诺", "质保", "维保"])
        if service and len(service) > 20:
            result["服务承诺"] = service[:800]

        return result

    def _extract_date_range(self, text: str, keywords: list[str]) -> tuple[str, str] | None:
        context = self._find_context(text, keywords, radius=250)
        if not context:
            return None
        dates = [self._format_date_match(m) for m in re.finditer(r"\d{4}\s*[年\-/\.]\s*\d{1,2}\s*[月\-/\.]\s*\d{1,2}\s*日?(?:\s*\d{1,2}\s*[：:时]\s*\d{1,2})?", context)]
        dates = [d for d in dates if d]
        if len(dates) >= 2:
            return dates[0], dates[1]
        return None

    def _extract_first_datetime(self, text: str) -> str:
        m = re.search(r"\d{4}\s*[年\-/\.]\s*\d{1,2}\s*[月\-/\.]\s*\d{1,2}\s*日?(?:\s*\d{1,2}\s*[：:时]\s*\d{1,2})?", text)
        return self._format_date_match(m) if m else ""

    def _format_date_match(self, match: re.Match[str] | None) -> str:
        if not match:
            return ""
        raw = match.group(0)
        m = re.search(r"(\d{4})\s*[年\-/\.]\s*(\d{1,2})\s*[月\-/\.]\s*(\d{1,2})", raw)
        if not m:
            return raw.strip()
        value = f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
        t = re.search(r"(\d{1,2})\s*[：:时]\s*(\d{1,2})", raw)
        if t:
            value += f" {int(t.group(1)):02d}:{int(t.group(2)):02d}:00"
        return value

    def _find_context(self, text: str, keywords: list[str], radius: int = 500) -> str:
        for keyword in keywords:
            idx = text.find(keyword)
            if idx >= 0:
                return text[max(0, idx - radius): idx + radius]
        return ""

    @staticmethod
    def _is_useful_value(value: Any) -> bool:
        if value is None:
            return False
        text = str(value).strip()
        return bool(text and text not in {"待补充", "无", "None", "null", "未能识别"})

    def _curate_text_for_llm(self, full_text: str, max_chars: int = 40000) -> str:
        """Curate text for LLM field extraction: include head, tail, and keyword-rich paragraphs."""
        if len(full_text) <= max_chars:
            return full_text

        # Strategy: head + keyword-rich middle + tail
        keywords = ["评分", "评标", "技术要求", "资质", "资格", "废标", "无效", "星标", "★",
                    "合同", "付款", "交付", "保证金", "保修", "质保", "关键条款", "实质性要求", "否决"]
        head_len = max_chars // 3
        tail_len = max_chars // 3
        middle_budget = max_chars - head_len - tail_len

        head = full_text[:head_len]
        tail = full_text[-tail_len:]

        important_middle: list[str] = []
        middle_text = full_text[head_len:-tail_len]
        for para in middle_text.split("\n"):
            if any(kw in para for kw in keywords):
                important_middle.append(para)
                if sum(len(p) for p in important_middle) >= middle_budget:
                    break

        curated = head + "\n\n...[中间内容筛选]...\n\n" + "\n".join(important_middle) + "\n\n...[末尾内容]...\n\n" + tail
        return curated[:max_chars]

    def _merge_tiny_chunks(self, chunks: list[dict[str, Any]]) -> list[dict[str, Any]]:
        """Merge chunks smaller than MIN_CHUNK_CHARS with adjacent chunks."""
        if not chunks:
            return chunks
        merged: list[dict[str, Any]] = [chunks[0].copy()]
        for chunk in chunks[1:]:
            last = merged[-1]
            if len(chunk["content"]) < MIN_CHUNK_CHARS:
                # Merge into previous
                last["content"] = last["content"] + "\n\n" + chunk["content"]
                last["end_page"] = chunk["end_page"]
                last["title"] = last["title"] if len(last["title"]) > len(chunk["title"]) else chunk["title"]
            elif len(last["content"]) < MIN_CHUNK_CHARS:
                # Previous was tiny, merge current into it
                last["content"] = last["content"] + "\n\n" + chunk["content"]
                last["end_page"] = chunk["end_page"]
            else:
                merged.append(chunk.copy())
        return merged

    def _create_placeholder_sections(
        self, db: Session, project_id: str, filename: str, clear_existing: bool = True
    ) -> list[ParsingSectionSummary]:
        """Fallback when no text could be extracted."""
        if clear_existing:
            db.query(ParsingSection).filter(ParsingSection.project_id == project_id).delete()
        created: list[ParsingSection] = []
        for name in BUSINESS_SECTIONS:
            created.append(self._add_section(
                db, project_id, name, "商务",
                f"【{name}】\n\n上传文件：{filename}\n\n（请在下方编辑器中填写内容）",
                name in BUSINESS_STAR, filename
            ))
        for name in TECH_SECTIONS:
            created.append(self._add_section(
                db, project_id, name, "技术",
                f"【{name}】\n\n上传文件：{filename}\n\n（请在下方编辑器中填写内容）",
                name in TECH_STAR, filename
            ))
        db.commit()
        return [ParsingSectionSummary.model_validate(s) for s in created]

    def simulate_parse(self, db: Session, project_id: str, filename: str) -> list[ParsingSectionSummary]:
        """Legacy simulation entrypoint; now delegates to placeholder creation."""
        return self._create_placeholder_sections(db, project_id, filename)

    # --- APIs used by generation service ---

    def project_fields_map(self, db: Session, project_id: str) -> dict[str, str]:
        """Return extracted fields as a flat dict for generation service."""
        project = db.query(Project).filter(Project.id == project_id).first()
        if project and project.extracted_fields:
            result: dict[str, str] = {}
            for item in project.extracted_fields:
                if isinstance(item, dict):
                    result[item.get("label", "")] = str(item.get("value", ""))
            return result
        # Fallback: try to build from parsing_sections content
        sections = db.query(ParsingSection).filter(ParsingSection.project_id == project_id).all()
        for s in sections:
            if s.section_name == "评分规则解析":
                return {"评分重点": s.content}
        return {}

    def get_project_context(self, db: Session, project_id: str) -> ParsingContext:
        """Return parsing context for generation service."""
        sections = (
            db.query(ParsingSection)
            .filter(ParsingSection.project_id == project_id)
            .order_by(ParsingSection.created_at.asc())
            .all()
        )
        parse_sections = [
            ParseSection(
                title=s.section_name,
                page="-",
                state="已解析" if s.content and len(s.content) > 20 else "待填充",
                source_text=s.content[:200] if s.content else "",
                source_file=s.source_file or "",
            )
            for s in sections
        ]
        # Build a source excerpt from the first few content sections
        excerpts: list[str] = []
        for s in sections:
            if s.section_type in ("评审", "内容") and s.content:
                excerpts.append(f"[{s.section_name}]\n{s.content[:500]}")
            if len("\n".join(excerpts)) > 3000:
                break
        return ParsingContext(
            parse_sections=parse_sections,
            source_excerpt="\n\n".join(excerpts),
        )


# Module-level singleton for imports
parsing_service = ParsingService()


# Keep module-level functions for backward compatibility with existing code
def get_or_create_project(db: Session, project_id: str) -> Project:
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        from fastapi import HTTPException, status
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="项目不存在")
    return project


def list_sections(db: Session, project_id: str) -> list[ParsingSectionSummary]:
    return parsing_service.list_sections(db, project_id)


def get_section_detail(db: Session, project_id: str, section_id: str) -> ParsingSectionDetail:
    return parsing_service.get_section_detail(db, project_id, section_id)


def update_section(
    db: Session, project_id: str, section_id: str, payload: ParsingSectionUpdateRequest
) -> ParsingSectionDetail:
    return parsing_service.update_section(db, project_id, section_id, payload)


def simulate_parse(db: Session, project_id: str, filename: str) -> list[ParsingSectionSummary]:
    return parsing_service.simulate_parse(db, project_id, filename)
