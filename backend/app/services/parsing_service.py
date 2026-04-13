import re
import tempfile
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi import HTTPException, status
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.models.project import Project
from app.models.project_document import ProjectDocument
from app.models.project_document_version import ProjectDocumentVersion
from app.models.project_extracted_field import ProjectExtractedField
from app.models.project_parse_section import ProjectParseSection
from app.schemas.parsing import (
    ParsedFileInfo,
    ProjectDocumentResponse,
    ProjectDocumentUploadResponse,
    ProjectDocumentVersionResponse,
    ProjectParsingContextResponse,
    ProjectParsingFieldUpdateRequest,
    ProjectParsingRunResponse,
)
from app.schemas.workspace import ExtractedField, ParseSection
from app.services.object_storage import object_storage_service
from app.services.workspace_service import get_extracted_fields, get_parse_sections


class ParsingService:
    def get_project_context(self, db: Session, project_id: str) -> ProjectParsingContextResponse:
        self._get_project(db, project_id)
        documents = self._list_document_rows(db, project_id)
        sections = self._list_project_sections(db, project_id)
        fields = self._list_project_fields(db, project_id)

        if not sections:
            sections = get_parse_sections(db)
        if not fields:
            fields = get_extracted_fields(db)

        source_excerpt = self._latest_source_excerpt(documents) or self._fallback_excerpt(db)
        return ProjectParsingContextResponse(
            project_id=project_id,
            documents=[self._to_document_response(db, doc) for doc in documents],
            parse_sections=sections,
            extracted_fields=fields,
            source_excerpt=source_excerpt,
        )

    def upload_project_document(
        self,
        db: Session,
        *,
        project_id: str,
        filename: str,
        file_bytes: bytes,
        document_type: str,
    ) -> ProjectDocumentUploadResponse:
        self._get_project(db, project_id)
        stored = object_storage_service.put_bytes(
            project_id=project_id,
            file_name=filename,
            file_bytes=file_bytes,
        )

        is_zip = filename.lower().endswith(".zip")
        all_sections: list[ParseSection] = []
        all_fields: list[ExtractedField] = []
        parsed_files: list[dict] = []
        combined_text_parts: list[str] = []

        if is_zip:
            import logging
            logger = logging.getLogger(__name__)
            zipfile = __import__("zipfile")

            def extract_from_zip(zf, prefix=""):
                nonlocal all_sections, parsed_files, combined_text_parts
                for zinfo in zf.infolist():
                    if zinfo.is_dir():
                        continue
                    if "__MACOSX" in zinfo.filename or zinfo.filename.startswith("."):
                        continue
                    zdata = zf.read(zinfo.filename)
                    zfilename = prefix + zinfo.filename if prefix else zinfo.filename
                    suffix = self._file_suffix(zfilename)
                    logger.info(f"[ZIP解析] 处理文件: {zfilename}, 类型: {suffix}")

                    if suffix == "zip":
                        try:
                            nested_zip = zipfile.ZipFile(BytesIO(zdata))
                            extract_from_zip(nested_zip, zfilename + " -> ")
                            continue
                        except Exception as nested_e:
                            logger.warning(f"[ZIP解析] 无法解压嵌套ZIP {zfilename}: {nested_e}")
                            parsed_files.append({
                                "file_name": zfilename,
                                "file_type": "zip",
                                "parse_status": "解压失败",
                                "document_type": "嵌套ZIP",
                                "section_found": "",
                                "word_count": 0,
                                "skip_reason": f"嵌套ZIP解压失败: {nested_e}",
                            })
                            continue

                    if suffix in {"txt", "md", "pdf", "docx", "xlsx", "xls"}:
                        ztext = self._extract_text_from_single_file(zfilename, zdata)
                        logger.info(f"[ZIP解析] 文件 {zfilename} 提取文本长度: {len(ztext) if ztext else 0}")
                        if ztext and len(ztext) > 50:
                            file_sections = self._build_sections(ztext, source_file=zfilename)
                            logger.info(f"[ZIP解析] 文件 {zfilename} 识别章节: {[s.title for s in file_sections]}")
                            if file_sections:
                                all_sections.extend(file_sections)
                            else:
                                fallback_sections = self._build_fallback_sections(ztext, source_file=zfilename)
                                all_sections.extend(fallback_sections)
                            combined_text_parts.append(f"\n\n=== {zfilename} ===\n{ztext[:3000]}")
                            parsed_files.append({
                                "file_name": zfilename,
                                "file_type": suffix,
                                "parse_status": "已解析",
                                "document_type": self._guess_document_type(zfilename),
                                "section_found": ", ".join([s.title for s in all_sections[-10:] if s.source_file == zfilename]) if all_sections else "",
                                "word_count": len(ztext),
                                "skip_reason": "",
                            })
                            logger.info(f"[ZIP解析] 文件 {zfilename} 已添加到解析清单")
                        else:
                            logger.warning(f"[ZIP解析] 文件 {zfilename} 文本长度不足50字符，跳过")
                            parsed_files.append({
                                "file_name": zfilename,
                                "file_type": suffix,
                                "parse_status": "跳过",
                                "document_type": self._guess_document_type(zfilename),
                                "section_found": "",
                                "word_count": len(ztext) if ztext else 0,
                                "skip_reason": "文本长度不足50字符",
                            })
                    else:
                        logger.info(f"[ZIP解析] 文件 {zfilename} 类型 {suffix} 不支持，跳过")
                        parsed_files.append({
                            "file_name": zfilename,
                            "file_type": suffix,
                            "parse_status": "不支持",
                            "document_type": "未知",
                            "section_found": "",
                            "word_count": 0,
                            "skip_reason": f"不支持的文件类型: {suffix}",
                        })

            try:
                with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
                    all_zip_files = [zi.filename for zi in zf.infolist() if not zi.is_dir() and "__MACOSX" not in zi.filename and not zi.filename.startswith(".")]
                    logger.info(f"[ZIP解析] 压缩包内共有 {len(all_zip_files)} 个文件: {all_zip_files}")
                    extract_from_zip(zf)
                    logger.info(f"[ZIP解析] 最终解析文件数: {len(parsed_files)}")
            except Exception as e:
                logger.error(f"[ZIP解析] 解析失败: {e}")
                raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"ZIP解析失败: {e}")
            
            combined_text = "\n".join(combined_text_parts)
            all_fields = self._build_fields(combined_text)
        else:
            text, parse_status = self._extract_text_from_file(filename=filename, file_bytes=file_bytes)
            if not text:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Unable to extract text from uploaded file.",
                )
            combined_text_parts.append(text)
            all_sections = self._build_sections(text, source_file=filename)
            if not all_sections:
                all_sections = self._build_fallback_sections(text, source_file=filename)
            all_fields = self._build_fields(text)
            parsed_files.append({
                "file_name": filename,
                "file_type": self._file_suffix(filename),
                "parse_status": parse_status,
                "document_type": document_type,
                "section_found": ", ".join([s.title for s in all_sections]) if all_sections else "",
                "word_count": len(text),
            })

        document = self._get_or_create_document(
            db,
            project_id=project_id,
            filename=filename,
            document_type=document_type,
        )
        version_no = self._next_version_no(db, document.id)
        combined_text = "\n".join(combined_text_parts)
        version = ProjectDocumentVersion(
            id=f"proj-ver-{uuid4().hex[:10]}",
            document_id=document.id,
            project_id=project_id,
            version_no=version_no,
            file_name=filename,
            file_type=self._file_suffix(filename),
            document_type=document_type.strip() or "招标文件",
            storage_backend=stored.backend,
            object_key=stored.object_key,
            file_size=stored.size,
            parse_status="已解析",
            source_text=combined_text[:64000],
        )
        db.add(version)

        document.file_name = filename
        document.file_type = self._file_suffix(filename)
        document.document_type = document_type.strip() or "招标文件"
        document.parse_status = "已解析"
        document.source_text = combined_text[:64000]
        db.flush()

        self._replace_project_parse_results(
            db,
            project_id=project_id,
            document_id=document.id,
            sections=all_sections,
            fields=all_fields,
        )

        db.commit()
        db.refresh(document)

        from app.schemas.parsing import ParsedFileInfo
        return ProjectDocumentUploadResponse(
            document=self._to_document_response(db, document),
            parse_sections=all_sections,
            extracted_fields=all_fields,
            source_excerpt=combined_text[:600],
            parsed_files=[ParsedFileInfo(**pf) for pf in parsed_files],
            debug_info={
                "total_files_in_zip": len(parsed_files),
                "parsed_count": len([pf for pf in parsed_files if pf.get("parse_status") == "已解析"]),
                "skipped_count": len([pf for pf in parsed_files if pf.get("skip_reason")]),
                "all_files_in_zip": [pf.get("file_name", "") for pf in parsed_files],
            },
        )

    def rerun_project_parsing(self, db: Session, project_id: str) -> ProjectParsingRunResponse:
        self._get_project(db, project_id)
        latest_version = db.scalars(
            select(ProjectDocumentVersion)
            .where(ProjectDocumentVersion.project_id == project_id)
            .order_by(ProjectDocumentVersion.created_at.desc())
            .limit(1)
        ).first()
        if latest_version is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"No project document found for '{project_id}'.",
            )

        sections = self._build_sections(latest_version.source_text)
        fields = self._build_fields(latest_version.source_text)
        self._replace_project_parse_results(
            db,
            project_id=project_id,
            document_id=latest_version.document_id,
            sections=sections,
            fields=fields,
        )
        db.commit()
        return ProjectParsingRunResponse(
            project_id=project_id,
            parse_sections=sections,
            extracted_fields=fields,
            source_excerpt=latest_version.source_text[:600],
        )

    def update_project_field(
        self,
        db: Session,
        project_id: str,
        field_label: str,
        payload: ProjectParsingFieldUpdateRequest,
    ) -> ExtractedField:
        self._get_project(db, project_id)
        row = db.scalars(
            select(ProjectExtractedField)
            .where(
                ProjectExtractedField.project_id == project_id,
                ProjectExtractedField.label == field_label,
            )
            .limit(1)
        ).first()
        if row is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Field '{field_label}' not found for project '{project_id}'.",
            )
        row.value = payload.value.strip()
        row.confidence = "人工校正"
        db.commit()
        db.refresh(row)
        return ExtractedField(label=row.label, value=row.value, confidence=row.confidence)

    def project_fields_map(self, db: Session, project_id: str) -> dict[str, str]:
        rows = db.scalars(
            select(ProjectExtractedField).where(ProjectExtractedField.project_id == project_id)
        ).all()
        return {row.label: row.value for row in rows}

    def _replace_project_parse_results(
        self,
        db: Session,
        *,
        project_id: str,
        document_id: str,
        sections: list[ParseSection],
        fields: list[ExtractedField],
    ) -> None:
        db.execute(delete(ProjectParseSection).where(ProjectParseSection.project_id == project_id))
        db.execute(delete(ProjectExtractedField).where(ProjectExtractedField.project_id == project_id))
        db.flush()

        for index, section in enumerate(sections, start=1):
            db.add(
                ProjectParseSection(
                    id=f"proj-sec-{uuid4().hex[:10]}",
                    project_id=project_id,
                    document_id=document_id,
                    title=section.title,
                    page=section.page,
                    state=section.state,
                    sort_order=index,
                    source_text=section.source_text or "",
                    source_file=section.source_file or "",
                )
            )
        for field in fields:
            db.add(
                ProjectExtractedField(
                    id=f"proj-field-{uuid4().hex[:10]}",
                    project_id=project_id,
                    document_id=document_id,
                    label=field.label,
                    value=field.value,
                    confidence=field.confidence,
                )
            )

    def _list_document_rows(self, db: Session, project_id: str) -> list[ProjectDocument]:
        return db.scalars(
            select(ProjectDocument)
            .where(ProjectDocument.project_id == project_id)
            .order_by(ProjectDocument.updated_at.desc())
        ).all()

    def _list_project_sections(self, db: Session, project_id: str) -> list[ParseSection]:
        rows = db.scalars(
            select(ProjectParseSection)
            .where(ProjectParseSection.project_id == project_id)
            .order_by(ProjectParseSection.sort_order.asc())
        ).all()
        return [ParseSection(title=row.title, page=row.page, state=row.state, source_text=row.source_text, source_file=row.source_file) for row in rows]

    def _list_project_fields(self, db: Session, project_id: str) -> list[ExtractedField]:
        rows = db.scalars(
            select(ProjectExtractedField)
            .where(ProjectExtractedField.project_id == project_id)
            .order_by(ProjectExtractedField.created_at.asc())
        ).all()
        return [ExtractedField(label=row.label, value=row.value, confidence=row.confidence) for row in rows]

    def _to_document_response(self, db: Session, row: ProjectDocument) -> ProjectDocumentResponse:
        versions = db.scalars(
            select(ProjectDocumentVersion)
            .where(ProjectDocumentVersion.document_id == row.id)
            .order_by(ProjectDocumentVersion.version_no.desc())
        ).all()
        return ProjectDocumentResponse(
            id=row.id,
            project_id=row.project_id,
            file_name=row.file_name,
            file_type=row.file_type,
            document_type=row.document_type,
            parse_status=row.parse_status,
            created_at=row.created_at,
            updated_at=row.updated_at,
            latest_version_no=versions[0].version_no if versions else 1,
            versions=[ProjectDocumentVersionResponse.model_validate(item) for item in versions],
        )

    def _latest_source_excerpt(self, documents: list[ProjectDocument]) -> str:
        if not documents:
            return ""
        return documents[0].source_text

    def _get_or_create_document(
        self,
        db: Session,
        *,
        project_id: str,
        filename: str,
        document_type: str,
    ) -> ProjectDocument:
        existing = db.scalars(
            select(ProjectDocument)
            .where(
                ProjectDocument.project_id == project_id,
                ProjectDocument.file_name == filename,
                ProjectDocument.document_type == (document_type.strip() or "招标文件"),
            )
            .limit(1)
        ).first()
        if existing is not None:
            return existing

        document = ProjectDocument(
            id=f"proj-doc-{uuid4().hex[:10]}",
            project_id=project_id,
            file_name=filename,
            file_type=self._file_suffix(filename),
            document_type=document_type.strip() or "招标文件",
            parse_status="待解析",
            source_text="",
        )
        db.add(document)
        db.flush()
        return document

    def _next_version_no(self, db: Session, document_id: str) -> int:
        versions = db.scalars(
            select(ProjectDocumentVersion)
            .where(ProjectDocumentVersion.document_id == document_id)
            .order_by(ProjectDocumentVersion.version_no.desc())
            .limit(1)
        ).all()
        return (versions[0].version_no + 1) if versions else 1

    def _build_sections(self, text: str, source_file: str = "") -> list[ParseSection]:
        sections: list[ParseSection] = []
        section_rules = [
            ("招标公告", ("招标公告", "项目概况", "招标范围")),
            ("资格审查条件", ("资格审查", "投标人资格", "资质要求")),
            ("评分办法", ("评分办法", "评分标准", "评审标准")),
            ("技术规范书", ("技术规范", "技术要求", "建设内容", "功能要求")),
            ("合同条款", ("合同条款", "付款条款", "违约责任", "服务承诺")),
        ]
        normalized = text.replace(" ", "")

        section_positions: list[tuple[str, str, str, int, int]] = []
        for idx, (title, keywords) in enumerate(section_rules, start=1):
            if any(keyword.replace(" ", "") in normalized for keyword in keywords):
                keyword_pos = -1
                for keyword in keywords:
                    pos = normalized.find(keyword.replace(" ", ""))
                    if pos != -1:
                        if keyword_pos == -1 or pos < keyword_pos:
                            keyword_pos = pos
                if keyword_pos != -1:
                    state = "高风险" if title == "合同条款" and "付款条款" in normalized else "已抽取"
                    section_positions.append((title, f"P{idx:02d}", state, keyword_pos, keyword_pos))

        section_positions.sort(key=lambda x: x[3])

        if len(section_positions) >= 2:
            for i in range(len(section_positions) - 1):
                title, page, state, start_pos, _ = section_positions[i]
                next_start = section_positions[i + 1][3]
                section_text = text[start_pos:next_start].strip()
                sections.append(ParseSection(title=title, page=page, state=state, source_text=section_text, source_file=source_file))
            last_title, last_page, last_state, last_start, _ = section_positions[-1]
            last_text = text[last_start:].strip()
            sections.append(ParseSection(title=last_title, page=last_page, state=last_state, source_text=last_text, source_file=source_file))
        elif len(section_positions) == 1:
            title, page, state, start_pos, _ = section_positions[0]
            section_text = text[start_pos:].strip()
            sections.append(ParseSection(title=title, page=page, state=state, source_text=section_text, source_file=source_file))
        return sections

    def _build_fallback_sections(self, text: str, source_file: str = "") -> list[ParseSection]:
        return [ParseSection(title="未分类全文", page="P01", state="待确认", source_text=text[:16000], source_file=source_file)]

    def _build_fields(self, text: str) -> list[ExtractedField]:
        # 先尝试使用 LLM 进行抽取
        try:
            from app.services.llm_parsing_client import llm_parsing_client
            llm_result = llm_parsing_client.extract_tender_fields(text)
            if llm_result and "项目名称" in llm_result:
                fields: list[ExtractedField] = []
                for label, info in llm_result.items():
                    if isinstance(info, dict):
                        value = self._normalize_field_value(label, str(info.get("value", ""))[:800])
                        confidence = str(info.get("confidence", "76%"))
                    else:
                        value = self._normalize_field_value(label, str(info)[:800])
                        confidence = "92%" if value != "待补充" else "76%"
                    fields.append(ExtractedField(label=label, value=value, confidence=confidence))
                return fields
        except Exception as e:
            print(f"LLM field extraction failed, falling back to regex: {e}")

        # 如果 LLM 失败，回退到正则匹配
        lines = [item.strip() for item in re.split(r"[\n\r]+", text) if item.strip()]
        joined = "\n".join(lines)
        scoring_focus = self._extract_scoring_focus(lines, joined)
        candidates = [
            ("项目名称", self._search_value(joined, [r"项目名称[:：]\s*(.+)", r"项目概况[:：]\s*(.+)"]) or self._first_nonempty(lines)),
            ("招标编号", self._search_value(joined, [r"(?:招标编号|项目编号)[:：]\s*(.+)"]) or "待补充"),
            ("投标截止时间", self._search_value(joined, [r"(?:投标截止时间|提交截止时间)[:：]\s*(.+)"]) or "待补充"),
            ("预算金额", self._search_value(joined, [r"(?:预算金额|最高限价)[:：]\s*(.+)"]) or "待补充"),
            ("必备资质", self._search_value(joined, [r"(?:资质要求|资格要求)[:：]\s*(.+)"]) or "待补充"),
            ("付款条款", self._search_value(joined, [r"(?:付款条款|付款方式)[:：]\s*(.+)"]) or "待补充"),
            ("交付周期", self._search_value(joined, [r"(?:交付周期|交货期|工期要求)[:：]\s*(.+)"]) or "待补充"),
            ("评分重点", scoring_focus or self._search_value(joined, [r"(?:评分办法|评分标准)[:：]\s*(.+)"]) or "待补充"),
            ("技术要求", self._search_value(joined, [r"(?:技术要求|功能要求|建设内容)[:：]\s*(.+)"]) or "待补充"),
            ("服务承诺", self._search_value(joined, [r"(?:服务承诺|售后服务|SLA)[:：]\s*(.+)"]) or "待补充"),
        ]
        fields: list[ExtractedField] = []
        for label, value in candidates:
            value = self._normalize_field_value(label, value[:800])
            confidence = "92%" if value != "待补充" else "76%"
            fields.append(ExtractedField(label=label, value=value, confidence=confidence))
        return fields

    def _extract_text_from_file(self, *, filename: str, file_bytes: bytes) -> tuple[str, str]:
        import zipfile

        def _parse_bytes(name: str, data: bytes) -> str:
            suffix = self._file_suffix(name)
            if suffix in {"txt", "md"}:
                return data.decode("utf-8", errors="ignore").strip()
            if suffix == "docx":
                try:
                    document = Document(BytesIO(data))
                    return "\n".join([para.text.strip() for para in document.paragraphs if para.text.strip()]).strip()
                except Exception as e:
                    return f"[解析DOCX失败 {name}: {e}]"
            if suffix == "pdf":
                try:
                    text = self._extract_pdf_text(data)
                    if self._looks_like_scanned_pdf(text):
                        ocr_text = self._ocr_pdf(data)
                        if ocr_text:
                            return ocr_text
                    return text
                except Exception as e:
                    return f"[解析PDF失败 {name}: {e}]"
            if suffix in {"xlsx", "xls"}:
                try:
                    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
                    text_parts = []
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        text_parts.append(f"--- Sheet: {sheet_name} ---")
                        rows_data = []
                        for row in ws.iter_rows(values_only=True):
                            row_values = [str(cell) if cell is not None else "" for cell in row]
                            if any(v.strip() for v in row_values):
                                rows_data.append("\t".join(row_values))
                        if rows_data:
                            text_parts.append("\n".join(rows_data))
                    wb.close()
                    return "\n".join(text_parts)
                except Exception as e:
                    return f"[解析Excel失败 {name}: {e}]"
            if suffix == "zip":
                try:
                    text_parts = []
                    with zipfile.ZipFile(BytesIO(data)) as zf:
                        for zinfo in zf.infolist():
                            if not zinfo.is_dir():
                                if "__MACOSX" in zinfo.filename or zinfo.filename.startswith("."):
                                    continue
                                zdata = zf.read(zinfo)
                                ztext = _parse_bytes(zinfo.filename, zdata)
                                if ztext:
                                    text_parts.append(f"\n\n=== 附件: {zinfo.filename} ===\n{ztext}")
                    return "".join(text_parts)
                except Exception as e:
                    return f"[解析ZIP失败 {name}: {e}]"
            return ""

        text = _parse_bytes(filename, file_bytes)
        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="Unsupported file type or unable to extract text. Please upload txt, pdf, docx, xlsx or zip.",
            )
        return text.strip(), "已解析"

    def _extract_text_from_single_file(self, filename: str, file_bytes: bytes) -> str:
        suffix = self._file_suffix(filename)
        if suffix in {"txt", "md"}:
            return file_bytes.decode("utf-8", errors="ignore").strip()
        if suffix == "docx":
            try:
                document = Document(BytesIO(file_bytes))
                return "\n".join([para.text.strip() for para in document.paragraphs if para.text.strip()]).strip()
            except Exception as e:
                return f"[解析DOCX失败 {filename}: {e}]"
        if suffix == "pdf":
            try:
                text = self._extract_pdf_text(file_bytes)
                if self._looks_like_scanned_pdf(text):
                    ocr_text = self._ocr_pdf(file_bytes)
                    if ocr_text:
                        return ocr_text
                return text
            except Exception as e:
                return f"[解析PDF失败 {filename}: {e}]"
        if suffix in {"xlsx", "xls"}:
            try:
                wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
                text_parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text_parts.append(f"--- Sheet: {sheet_name} ---")
                    rows_data = []
                    for row in ws.iter_rows(values_only=True):
                        row_values = [str(cell) if cell is not None else "" for cell in row]
                        if any(v.strip() for v in row_values):
                            rows_data.append("\t".join(row_values))
                    if rows_data:
                        text_parts.append("\n".join(rows_data))
                wb.close()
                return "\n".join(text_parts)
            except Exception as e:
                return f"[解析Excel失败 {filename}: {e}]"
        return ""

    def _guess_document_type(self, filename: str) -> str:
        fname_lower = filename.lower()
        type_mapping = {
            "招标公告": ["招标公告", "公告"],
            "资格审查条件": ["资格审查", "资质", "资格预审"],
            "评分办法": ["评分", "评审", "评标"],
            "技术规范书": ["技术", "规范", "技术规范"],
            "合同条款": ["合同", "条款", "协议书"],
        }
        for doc_type, keywords in type_mapping.items():
            for kw in keywords:
                if kw in fname_lower:
                    return doc_type
        return "招标文件"

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        reader = PdfReader(BytesIO(file_bytes))
        page_texts: list[str] = []
        for page in reader.pages:
            extracted = (page.extract_text() or "").strip()
            if extracted:
                page_texts.append(extracted)
        return "\n".join(page_texts).strip()

    def _looks_like_scanned_pdf(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", text)
        return len(normalized) < 80

    def _ocr_pdf(self, file_bytes: bytes) -> str:
        texts: list[str] = []
        image_paths = self._render_pdf_to_images(file_bytes)
        if not image_paths:
            return ""
        try:
            for image_path in image_paths[:8]:
                page_text = self._ocr_image(Path(image_path))
                if page_text:
                    texts.append(page_text.strip())
            return "\n".join([item for item in texts if item]).strip()
        finally:
            for image_path in image_paths:
                try:
                    Path(image_path).unlink(missing_ok=True)
                except Exception:
                    pass

    def _render_pdf_to_images(self, file_bytes: bytes) -> list[str]:
        try:
            import fitz
        except ImportError:
            return []

        temp_paths: list[str] = []
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            for index, page in enumerate(pdf):
                pixmap = page.get_pixmap(dpi=180)
                tmp = tempfile.NamedTemporaryFile(prefix=f"ocr-page-{index}-", suffix=".png", delete=False)
                tmp.close()
                pixmap.save(tmp.name)
                temp_paths.append(tmp.name)
        return temp_paths

    def _ocr_image(self, image_path: Path) -> str:
        try:
            from rapidocr_onnxruntime import RapidOCR

            engine = RapidOCR()
            result, _elapsed = engine(str(image_path))
            if result:
                text = "\n".join([item[1] for item in result if len(item) > 1 and item[1].strip()])
                if text.strip():
                    return text.strip()
        except Exception:
            pass

        try:
            from ocrmac.ocrmac import text_from_image

            annotations = text_from_image(
                str(image_path),
                recognition_level="accurate",
                language_preference=["zh-Hans", "en-US"],
                confidence_threshold=0.0,
                detail=True,
            )
            if annotations:
                text = "\n".join([item[0] for item in annotations if item and item[0].strip()])
                if text.strip():
                    return text.strip()
        except Exception:
            pass

        try:
            import pytesseract
            from PIL import Image

            with Image.open(image_path) as image:
                return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
        except Exception:
            return ""

    def _search_value(self, text: str, patterns: list[str]) -> str:
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()[:800]
        return ""

    def _normalize_field_value(self, label: str, value: str) -> str:
        cleaned = re.sub(r"[ \t]+", " ", value).strip()
        if label == "投标截止时间":
            cleaned = re.sub(r"(\d{4}-\d{2}-\d{2})(\d{2}:\d{2})", r"\1 \2", cleaned)
            cleaned = re.sub(r"(\d{4}/\d{2}/\d{2})(\d{2}:\d{2})", r"\1 \2", cleaned)
        return cleaned

    def _first_nonempty(self, lines: list[str]) -> str:
        for line in lines:
            if len(line) >= 4:
                return line[:120]
        return "待补充"

    def _extract_scoring_focus(self, lines: list[str], text: str) -> str:
        scoring_lines: list[str] = []
        for line in lines:
            normalized = re.sub(r"\s+", "", line)
            has_score = bool(re.search(r"(\d+(?:\.\d+)?)\s*分", line))
            has_keyword = any(
                keyword in normalized
                for keyword in ("评分", "评审", "分值", "技术", "商务", "业绩", "服务", "资质", "团队", "案例")
            )
            if has_score and has_keyword:
                scoring_lines.append(line[:160])
        if not scoring_lines:
            block_match = re.search(
                r"(?:评分办法|评分标准|评审标准)(.*?)(?:技术规范|合同条款|投标截止|$)",
                text,
                flags=re.DOTALL,
            )
            if block_match:
                block_lines = [
                    item.strip()
                    for item in re.split(r"[\n\r]+", block_match.group(1))
                    if item.strip()
                ]
                for line in block_lines:
                    if re.search(r"(\d+(?:\.\d+)?)\s*分", line):
                        scoring_lines.append(line[:160])
        dedup: list[str] = []
        seen: set[str] = set()
        for line in scoring_lines:
            key = re.sub(r"\s+", "", line)
            if key in seen:
                continue
            seen.add(key)
            dedup.append(line)
        return "；".join(dedup[:8])

    def _fallback_excerpt(self, db: Session) -> str:
        fields = get_extracted_fields(db)
        return "\n".join([f"{field.label}：{field.value}" for field in fields[:4]])

    def _get_project(self, db: Session, project_id: str) -> Project:
        project = db.get(Project, project_id)
        if project is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Project '{project_id}' not found.",
            )
        return project

    def _file_suffix(self, filename: str) -> str:
        return filename.lower().rsplit(".", 1)[-1] if "." in filename else ""


parsing_service = ParsingService()
